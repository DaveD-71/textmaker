"""
Post-process a DOCX produced by Pandoc to ensure:
- (opt-in, --toc) a section break is inserted immediately after the Table of Contents,
  page numbering restarts at 1, and page numbers are removed from TOC pages
- (opt-in, --h1-sections) section breaks (nextPage) are inserted before each H1 heading;
  omit this when the reference DOCX already sets pageBreakBefore on Heading 1
- Quick Parts are inserted through Word COM from a real template when available

Semantic label rendering (icon tables, character styles, unit title tables) is opt-in;
pass --apply-semantic-labels to enable it. Style definitions are never created or
redefined here — all style management belongs in manage_docx_styles.py.
"""
# pylint: disable=protected-access,broad-exception-caught
from __future__ import annotations

import argparse
import copy
import csv
import ctypes
import os
import re
import shutil
import subprocess
import sys
import tempfile
import threading
import time
from contextlib import contextmanager
from pathlib import Path
from typing import Iterable

try:
    from docx import Document  # type: ignore[reportMissingImports]
    from docx.enum.style import WD_STYLE_TYPE  # type: ignore[reportMissingImports]
    from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore[reportMissingImports]
    from docx.oxml import OxmlElement  # type: ignore[reportMissingImports]
    from docx.oxml.ns import qn  # type: ignore[reportMissingImports]
    from docx.oxml.table import CT_Tbl  # type: ignore[reportMissingImports]
    from docx.oxml.text.paragraph import CT_P  # type: ignore[reportMissingImports]
    from docx.table import Table  # type: ignore[reportMissingImports]
    from docx.text.paragraph import Paragraph  # type: ignore[reportMissingImports]
except ImportError as exc:
    raise RuntimeError(
        'Missing dependency: python-docx is required. Install with `pip install python-docx`.'
    ) from exc

try:
    import pythoncom  # type: ignore[reportMissingImports]
    import win32com.client  # type: ignore[reportMissingImports]
except ImportError:
    pythoncom = None
    win32com = None


NS = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

WHY_LABEL = 'Why this works:'
CHECK_LABELS = ('Before you write:',)
LEARN_LABELS = (
    'Teaching point:',
    'Key principle:',
    'Writing goal:',
    'Transfer reminder:',
    'Role reminder:',
    'Planning reminder:',
)
NOTE_LABELS = (
    'Note:',
    'Note on',
    'Additional note:',
    'How to use the checklist:',
    'Rubric note:',
    'Self-study note:',
)
BAD_MODEL_HEADINGS = {
    'Original Version',
    'Original Version (too direct)',
    'Original (Disjointed)',
}
GOOD_MODEL_HEADINGS = {
    'Revised Version',
    'Revised Version (diplomatic)',
    'Revised (Cohesive)',
}
NEUTRAL_MODEL_SOURCE_STYLES = {
    'Block Text',
    'Quote',
    'Intense Quote',
}
QUOTED_MODEL_RE = re.compile(r'^[\u201c"].+[\u201d"]$')
WORD_COUNT_RE = re.compile(
    r'\b(?:approximately|about|around)?\s*\d+\s*(?:[-\u2013\u2014]\s*\d+)?\s*words?\b',
    re.IGNORECASE,
)
NO_TITLE_MARKER_RE = re.compile(r'^\s*no\s+title\s*$', re.IGNORECASE)
UNIT_HEADING_RE = re.compile(r'^Unit\s+(\d+)(?:\s+[-\u2013\u2014]\s+|\.\s+)(.+)$')
MODULE_HEADING_RE = re.compile(r'^Module\s+\d+(?:\s+[-\u2013\u2014]\s+|\.\s+).+$', re.IGNORECASE)
MODULE_REVIEW_HEADING_RE = re.compile(r'^Module\s+\d+\s+Review\s+Workshop$', re.IGNORECASE)
ALPHA_ORDINAL_RE = re.compile(r'^([A-Z])[\.)]\s+\S+')
ACTIVITY_CODE_SUFFIX_RE = re.compile(r'\s+\(([A-H]\d)(?:[^)]*)\)\s*(?:[★*])?\s*$')
MODULE_UNIT_RANGE_SUFFIX_RE = re.compile(
    r'\s+\(Units?\s+\d+\s*[-\u2013\u2014]\s*\d+\)\s*$',
    re.IGNORECASE,
)


# Maps paragraph style ID → tag_filled_*.png filename stem.
# example-good and example-bad share the example tag image.
DIV_TAG_ICON_STEMS = {
    'DivLabelLearn':       'tag_filled_learn',
    'DivLabelLanguage':    'tag_filled_language',
    'DivLabelStructure':   'tag_filled_structure',
    'DivLabelNotice':      'tag_filled_notice',
    'DivLabelWrite':       'tag_filled_write',
    'DivLabelRewrite':     'tag_filled_rewrite',
    'DivLabelRevise':      'tag_filled_revise',
    'DivLabelEdit':        'tag_filled_edit',
}

# Height (in EMU) at which to render the tag icon inline.
# 198000 EMU = 0.55 cm
DIV_TAG_ICON_HEIGHT_EMU = 198000

# distR (EMU) gap to the right of the inline image before label text.
# 114300 EMU = 9pt
DIV_TAG_ICON_DIST_R_EMU = 114300

# Which tag row to use: 'filled' or 'outline'.
DIV_TAG_ICON_STYLE = 'filled'


def _resolve_div_tag_icon(
    style_id: str,
    reference_doc_path,
    tag_style: str = DIV_TAG_ICON_STYLE,
) -> 'Path | None':
    """Return the Path to the tag icon PNG for style_id, or None if not found.

    Looks for div-tags-icons-2_assets/ as a sibling of the reference DOCX file.
    tag_style is 'filled' or 'outline'.
    """
    stem = DIV_TAG_ICON_STEMS.get(style_id)
    if not stem or not reference_doc_path:
        return None
    # Replace 'tag_filled_' prefix with the requested style variant.
    stem = stem.replace('tag_filled_', f'tag_{tag_style}_')
    ref_dir = Path(reference_doc_path).parent
    icon_path = ref_dir / 'div-tags-icons-2_assets' / f'{stem}.png'
    return icon_path if icon_path.exists() else None


def build_semantic_div_label_styles(reference_doc_path=None, context: PostprocessContext | None = None):
    """
    Build a mapping from paragraph style ID to hex color string for semantic Div label runs.

    Reads color directly from the paragraph style's rPr — no linked char styles required.
    Keys are paragraph style IDs (e.g. 'DivLabelNotice'); values are hex color strings (e.g. 'DB4351').
    """
    if context is not None:
        return context.semantic_div_label_styles()
    if not reference_doc_path:
        return {}

    ref_doc = Document(reference_doc_path)
    mapping = {}
    for style in ref_doc.styles:
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if not style.name.startswith('Div Label '):
            continue
        color = _style_color_value(style)
        if color:
            mapping[style.style_id] = color
    return mapping


# SEMANTIC_DIV_LABEL_STYLES will be set in apply_semantic_div_labels

LEARN_SEMANTIC_STYLE_IDS: set[str] = set()  # Learn base conversion removed


DEFAULT_BUILDING_BLOCK_TEMPLATE = (
    Path(os.environ.get('APPDATA', str(Path.home() / 'AppData' / 'Roaming')))
    / 'Microsoft'
    / 'Templates'
    / 'Normal.dotm'
)
BUILDING_BLOCK_NAMES = {'unit': ('Call Out - Unit Title', 'Unit Tile')}

WORD_COLLAPSE_START = 1
WORD_COLLAPSE_END = 0
WORD_PAGE_BREAK = 7
WORD_ALERTS_NONE = 0
WORD_DO_NOT_SAVE_CHANGES = 0
WORD_SAVE_CHANGES = -1


class PostprocessProfiler:
    """Progress logger with optional timing summary for DOCX postprocess passes."""

    def __init__(
        self,
        enabled: bool = False,
        *,
        progress: bool = True,
        warn_after_seconds: float = 30.0,
    ):
        self.enabled = enabled
        self.progress = progress
        self.warn_after_seconds = max(0.0, warn_after_seconds)
        self.records: list[tuple[str, float, int | None, str]] = []

    @contextmanager
    def step(self, name: str, *, changed: int | None = None, status: str = 'done'):
        if self.progress:
            print(f'[progress] postprocess:{name}: started', flush=True)
        start = time.perf_counter()
        result = {'changed': changed, 'status': status}
        stop_event = threading.Event()
        watchdog = self._start_watchdog(name, start, stop_event)
        try:
            yield result
        except Exception:
            result['status'] = 'error'
            elapsed = time.perf_counter() - start
            if self.progress:
                print(
                    f'[error] postprocess:{name}: failed after {elapsed:.3f}s',
                    flush=True,
                )
            raise
        finally:
            stop_event.set()
            if watchdog is not None:
                watchdog.join(timeout=0.2)
            elapsed = time.perf_counter() - start
            if self.progress:
                changed_value = result['changed']
                status_value = result['status']
                count_text = '' if changed_value is None else f' count={changed_value}'
                print(
                    f'[progress] postprocess:{name}: {status_value} in {elapsed:.3f}s{count_text}',
                    flush=True,
                )
            if self.enabled:
                changed_value = result['changed']
                status_value = result['status']
                self.records.append((name, elapsed, changed_value, status_value))
                count_text = '' if changed_value is None else f' count={changed_value}'
                print(f'[profile] postprocess:{name}: {elapsed:.3f}s status={status_value}{count_text}')

    def skip(self, name: str) -> None:
        if self.progress:
            print(f'[progress] postprocess:{name}: skipped', flush=True)
        if self.enabled:
            self.records.append((name, 0.0, 0, 'skipped'))
            print(f'[profile] postprocess:{name}: 0.000s skipped changed=0')

    def _start_watchdog(self, name: str, start: float, stop_event: threading.Event):
        if not self.progress or self.warn_after_seconds <= 0:
            return None

        def watch() -> None:
            while not stop_event.wait(self.warn_after_seconds):
                elapsed = time.perf_counter() - start
                print(
                    f'[warn] postprocess:{name}: still running after {elapsed:.1f}s',
                    flush=True,
                )

        thread = threading.Thread(target=watch, daemon=True)
        thread.start()
        return thread

    def summary(self) -> None:
        if not self.enabled or not self.records:
            return
        total = sum(elapsed for _, elapsed, _, status in self.records if status != 'skipped')
        print('[profile] postprocess summary:')
        for name, elapsed, changed, status in self.records:
            count_text = '' if changed is None else f' count={changed}'
            print(f'[profile]   {name}: {elapsed:.3f}s status={status}{count_text}')
        print(f'[profile]   total measured: {total:.3f}s')


class PostprocessContext:
    """Cached resources shared across postprocess passes."""

    def __init__(self, reference_doc_path=None):
        self.reference_doc_path = Path(reference_doc_path) if reference_doc_path else None
        self._reference_doc = None
        self._reference_style_ids: set[str] | None = None
        self._semantic_div_label_styles: dict[str, str] | None = None
        self._unit_tile_prototype = None

    def reference_doc(self):
        if not self.reference_doc_path or not self.reference_doc_path.exists():
            return None
        if self._reference_doc is None:
            self._reference_doc = Document(self.reference_doc_path)
        return self._reference_doc

    def reference_style_ids(self) -> set[str]:
        if self._reference_style_ids is None:
            ref_doc = self.reference_doc()
            self._reference_style_ids = _style_ids(ref_doc.styles) if ref_doc is not None else set()
        return self._reference_style_ids

    def semantic_div_label_styles(self) -> dict[str, str]:
        if self._semantic_div_label_styles is None:
            mapping = {}
            ref_doc = self.reference_doc()
            if ref_doc is not None:
                for style in ref_doc.styles:
                    if style.type != WD_STYLE_TYPE.PARAGRAPH:
                        continue
                    if not style.name.startswith('Div Label '):
                        continue
                    color = _style_color_value(style)
                    if color:
                        mapping[style.style_id] = color
            self._semantic_div_label_styles = mapping
        return self._semantic_div_label_styles

    def unit_tile_prototype(self):
        if self._unit_tile_prototype is not None:
            return copy.deepcopy(self._unit_tile_prototype)
        ref_doc = self.reference_doc()
        if ref_doc is None:
            return None
        for table in ref_doc.tables:
            try:
                first = _clean_word_text(table.cell(0, 0).text)
                second = _clean_word_text(table.cell(0, 1).text)
            except Exception:
                continue
            if re.match(r'^(?:U0|U#|#|\d+)$', first) and 'Unit Title' in second:
                self._unit_tile_prototype = copy.deepcopy(table._tbl)
                return copy.deepcopy(self._unit_tile_prototype)
        return None


def _copy_page_layout_and_footers(source_sect, target_sect, *, copy_footers=False):
    """Copy stable section properties needed by generated section breaks.

    Pandoc/reference DOCX output usually keeps page-number fields in footer
    parts referenced from the document-level ``sectPr``. New section breaks
    need their own footer references after the postprocessor unlinks and
    rewrites headers section-by-section; otherwise Word can treat later
    sections as having no active footer and page numbers disappear.
    """
    if source_sect is None:
        return
    if copy_footers:
        insert_at = 0
        for ref in source_sect.findall(qn('w:footerReference')):
            ref_type = ref.get(qn('w:type'))
            # Do not propagate a first-page footer unless titlePg is also
            # deliberately copied. In this book the first footer is used to
            # suppress page numbers on front/TOC pages; copying it into every
            # generated section would hide first-page numbers throughout.
            if ref_type == 'first':
                continue
            target_sect.insert(insert_at, copy.deepcopy(ref))
            insert_at += 1
    for tag in (qn('w:pgSz'), qn('w:pgMar')):
        el = source_sect.find(tag)
        if el is not None:
            target_sect.append(copy.deepcopy(el))


def ensure_section_footer_references(doc, *, skip_first_section=False) -> int:
    """Ensure generated sections retain normal odd/even footer page numbers.

    Re-running the postprocessor or un-linking headers can leave many section
    properties with explicit header references but no footer references. Word
    then renders those sections without the reference footer PAGE field. This
    pass backfills the document's normal odd/even footer references into any
    section that lacks them.
    """
    body = doc.element.body
    sect_prs = list(body.findall('.//' + qn('w:sectPr')))
    if not sect_prs:
        return 0

    template_refs = []
    for sect_pr in reversed(sect_prs):
        refs = [
            ref for ref in sect_pr.findall(qn('w:footerReference'))
            if ref.get(qn('w:type')) in {'default', 'even'}
        ]
        if refs:
            template_refs = refs
            break
    if not template_refs:
        return 0

    changed = 0
    start_index = 1 if skip_first_section else 0
    for sect_pr in sect_prs[start_index:]:
        existing_types = {
            ref.get(qn('w:type'))
            for ref in sect_pr.findall(qn('w:footerReference'))
        }
        insert_at = 0
        for ref in template_refs:
            ref_type = ref.get(qn('w:type'))
            if ref_type in existing_types:
                continue
            sect_pr.insert(insert_at, copy.deepcopy(ref))
            insert_at += 1
            existing_types.add(ref_type)
            changed += 1
    return changed


def _insert_section_break_before_paragraph(
    paragraph,
    restart_numbering=False,
    start_page=1,
    copy_footers=True,
):
    """
    Insert a nextPage section break BEFORE a paragraph by attaching sectPr
    to the PREVIOUS paragraph (that's how Word section breaks work).

    If restart_numbering=True, the new section will restart page numbering at start_page.
    """
    doc_element = paragraph._element.getparent()
    paragraphs = list(doc_element.iterchildren())

    try:
        current_index = paragraphs.index(paragraph._element)
    except ValueError:
        return False

    if current_index == 0:
        return False

    prev_p = None
    for candidate in reversed(paragraphs[:current_index]):
        if candidate.tag == qn('w:p'):
            prev_p = candidate
            break
    if prev_p is None:
        return False

    p_pr = prev_p.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        prev_p.insert(0, p_pr)

    for sect in list(p_pr.findall(qn('w:sectPr'))):
        try:
            p_pr.remove(sect)
        except Exception:
            pass

    sect_pr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sect_pr.append(type_el)

    if restart_numbering:
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), str(start_page))
        sect_pr.append(pg_num_type)

    # Copy page size/margins and normal footer references from document-level
    # sectPr so inserted sections inherit the correct paper size and page
    # numbering instead of defaulting to a footerless section.
    doc_sect = doc_element.find(qn('w:sectPr')) if doc_element is not None else None
    _copy_page_layout_and_footers(doc_sect, sect_pr, copy_footers=copy_footers)

    p_pr.append(sect_pr)
    return True


def _apply_next_page_section_to_paragraph(paragraph, start_page=None, copy_footers=False):
    """Attach a nextPage sectPr (and optional page-number restart) to an existing paragraph.

    Copies w:pgSz and w:pgMar from the document's main sectPr so the inserted
    section break never falls back to the application default page size.
    """
    p_pr = paragraph._p.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        paragraph._p.insert(0, p_pr)

    for sect in list(p_pr.findall(qn('w:sectPr'))):
        try:
            p_pr.remove(sect)
        except Exception:
            pass

    sect_pr = OxmlElement('w:sectPr')
    type_el = OxmlElement('w:type')
    type_el.set(qn('w:val'), 'nextPage')
    sect_pr.append(type_el)

    # Copy page size and margins from the document-level sectPr. Footer
    # references are intentionally optional because the TOC/front-matter
    # section may need page numbers suppressed without mutating the shared
    # body footer parts.
    body = paragraph._p.getparent()
    doc_sect = body.find(qn('w:sectPr')) if body is not None else None
    _copy_page_layout_and_footers(doc_sect, sect_pr, copy_footers=copy_footers)

    if start_page is not None:
        pg_num_type = OxmlElement('w:pgNumType')
        pg_num_type.set(qn('w:start'), str(start_page))
        sect_pr.append(pg_num_type)

    p_pr.append(sect_pr)


def _remove_page_numbers_from_footer(footer):
    """Remove PAGE field codes from a footer to suppress page numbers."""
    try:
        for para in footer.paragraphs:
            runs_to_remove = []
            for run in para.runs:
                run_elem = run._element
                fld_simple = run_elem.find('.//w:fldSimple', NS)
                fld_char = run_elem.find('.//w:fldChar', NS)
                instr_text = run_elem.find('.//w:instrText', NS)

                if fld_simple is not None:
                    instr = fld_simple.get(qn('w:instr'), '')
                    if 'PAGE' in instr.upper():
                        runs_to_remove.append(run)
                elif instr_text is not None:
                    if 'PAGE' in (instr_text.text or '').upper():
                        runs_to_remove.append(run)
                elif fld_char is not None:
                    all_instr = run_elem.findall('.//w:instrText', NS)
                    for it in all_instr:
                        if 'PAGE' in (it.text or '').upper():
                            runs_to_remove.append(run)
                            break

            for run in runs_to_remove:
                para._element.remove(run._element)
    except Exception as exc:
        print(f'Warning: Could not remove page numbers from footer: {exc}')


def _is_heading_1(paragraph):
    """Check if paragraph is styled as Heading 1."""
    style = paragraph.style
    if style and 'Heading 1' in style.name:
        return True
    if hasattr(style, 'style_id') and style.style_id == 'Heading1':
        return True
    return False


def insert_section_breaks_before_h1(doc, skip_first=True, restart_first=False):
    """Insert nextPage section breaks before each H1 heading."""
    paragraphs = list(doc.paragraphs)
    h1_count = 0
    sections_added = 0

    for para in paragraphs:
        if _is_heading_1(para):
            h1_count += 1
            if skip_first and h1_count == 1:
                continue

            restart = restart_first and h1_count == 1
            try:
                success = _insert_section_break_before_paragraph(
                    para,
                    restart_numbering=restart,
                    start_page=1,
                )
                if success:
                    sections_added += 1
            except Exception as exc:
                print(f"Warning: Could not insert section break before H1 '{para.text}': {exc}")

    return sections_added


def _qn_attr(nsmap, attr):
    return f'{{{nsmap.get("w")}}}{attr}'


def _build_num_format_map(doc):
    """Map numId/ilvl to numFmt (e.g., bullet, decimal)."""
    try:
        num_part = doc.part.numbering_part.element
    except Exception:
        return {}
    nsmap = num_part.nsmap

    abstract_map = {}
    for abs_num in num_part.findall('w:abstractNum', nsmap):
        abs_id = abs_num.get(_qn_attr(nsmap, 'abstractNumId'))
        lvl_map = {}
        for lvl in abs_num.findall('w:lvl', nsmap):
            ilvl = lvl.get(_qn_attr(nsmap, 'ilvl'))
            numfmt_el = lvl.find('w:numFmt', nsmap)
            if ilvl is not None and numfmt_el is not None:
                lvl_map[ilvl] = numfmt_el.get(_qn_attr(nsmap, 'val'))
        if abs_id is not None:
            abstract_map[abs_id] = lvl_map

    num_map = {}
    for num in num_part.findall('w:num', nsmap):
        num_id = num.get(_qn_attr(nsmap, 'numId'))
        abs_el = num.find('w:abstractNumId', nsmap)
        if num_id is None or abs_el is None:
            continue
        abs_id = abs_el.get(_qn_attr(nsmap, 'val'))
        num_map[num_id] = abstract_map.get(abs_id, {})

    return num_map


def _numbering_element(doc):
    try:
        return doc.part.numbering_part.element
    except Exception:
        return None


def _num_id_from_style(style) -> str | None:
    style_el = getattr(style, '_element', None)
    if style_el is None:
        return None
    num_id_el = style_el.find('.//w:numPr/w:numId', NS)
    return num_id_el.get(qn('w:val')) if num_id_el is not None else None


def _paragraph_num_id(paragraph) -> str | None:
    p_pr = paragraph._p.find(qn('w:pPr'))
    num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None
    if num_pr is None:
        return None
    num_id_el = num_pr.find(qn('w:numId'))
    return num_id_el.get(qn('w:val')) if num_id_el is not None else None


def _abstract_num_id_for_num(numbering, num_id: str) -> str | None:
    nsmap = numbering.nsmap
    for num in numbering.findall('w:num', nsmap):
        if num.get(_qn_attr(nsmap, 'numId')) != str(num_id):
            continue
        abstract_el = num.find('w:abstractNumId', nsmap)
        if abstract_el is not None:
            return abstract_el.get(_qn_attr(nsmap, 'val'))
    return None


def _next_num_id(numbering) -> str:
    nsmap = numbering.nsmap
    used = []
    for num in numbering.findall('w:num', nsmap):
        num_id = num.get(_qn_attr(nsmap, 'numId'))
        if num_id and num_id.isdigit():
            used.append(int(num_id))
    return str((max(used) if used else 0) + 1)


def _create_restarted_num_id(doc, base_num_id: str, start_value: int, ilvl: str = '0') -> str | None:
    numbering = _numbering_element(doc)
    if numbering is None:
        return None
    abstract_num_id = _abstract_num_id_for_num(numbering, base_num_id)
    if abstract_num_id is None:
        return None

    new_num_id = _next_num_id(numbering)
    num = OxmlElement('w:num')
    num.set(qn('w:numId'), new_num_id)

    abstract = OxmlElement('w:abstractNumId')
    abstract.set(qn('w:val'), abstract_num_id)
    num.append(abstract)

    lvl_override = OxmlElement('w:lvlOverride')
    lvl_override.set(qn('w:ilvl'), ilvl)
    start_override = OxmlElement('w:startOverride')
    start_override.set(qn('w:val'), str(start_value))
    lvl_override.append(start_override)
    num.append(lvl_override)

    numbering.append(num)
    return new_num_id


def _apply_direct_num_id(paragraph, num_id: str, ilvl: str = '0') -> bool:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn('w:numPr'))
    if existing is not None:
        p_pr.remove(existing)

    num_pr = OxmlElement('w:numPr')
    ilvl_el = OxmlElement('w:ilvl')
    ilvl_el.set(qn('w:val'), ilvl)
    num_id_el = OxmlElement('w:numId')
    num_id_el.set(qn('w:val'), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    return True


def _alpha_marker_start_value(marker: str) -> int:
    return ord(marker.upper()) - ord('A') + 1


def _get_style_by_name_or_id(styles, target, style_type=None):
    """Fetch a style by name or style_id; returns None if missing."""
    try:
        style = styles[target]
        if style_type is None or getattr(style, 'type', None) == style_type:
            return style
    except Exception:
        pass
    for st in styles:
        if getattr(st, 'name', None) == target or getattr(st, 'style_id', None) == target:
            if style_type is not None and getattr(st, 'type', None) != style_type:
                continue
            return st
    return None


def _require_style(styles, target):
    style = _get_style_by_name_or_id(styles, target)
    if not style:
        raise RuntimeError(f'Required style is missing from the reference DOCX: {target}')
    return style


def _ensure_styles_from_reference(doc, reference_doc_path, style_names):
    """Copy named styles from reference DOCX into doc if they are not already present.

    Copies the raw <w:style> XML element from the reference DOCX styles part directly
    into the output DOCX styles part.  Only copies styles that are absent from doc.
    Returns the list of style names that were actually copied.
    """
    if not reference_doc_path:
        return []
    missing = [n for n in style_names if _get_style_by_name_or_id(doc.styles, n) is None]
    if not missing:
        return []

    ref_doc = Document(reference_doc_path)
    ref_styles_el = ref_doc.styles.element   # <w:styles> element
    out_styles_el = doc.styles.element

    copied = []
    for name in missing:
        ref_style = _get_style_by_name_or_id(ref_doc.styles, name)
        if ref_style is None:
            continue
        # Deep-copy the <w:style> element and append to output styles part
        out_styles_el.append(copy.deepcopy(ref_style.element))
        copied.append(name)

    return copied


def _require_character_style(styles, target):
    style = _get_style_by_name_or_id(styles, target, WD_STYLE_TYPE.CHARACTER)
    if not style:
        raise RuntimeError(f'Required character style is missing from the reference DOCX: {target}')
    return style


def _normalize_text(text: str) -> str:
    return ' '.join(text.split())


def _clean_word_text(text: str) -> str:
    return _normalize_text(text.replace('\r', ' ').replace('\x07', ' '))


def _starts_with_any(text: str, prefixes: Iterable[str]) -> bool:
    return any(text.startswith(prefix) for prefix in prefixes)


def _is_heading(paragraph) -> bool:
    style = paragraph.style
    return bool(style and getattr(style, 'name', '').startswith('Heading'))


def _is_paragraph_italic(paragraph) -> bool:
    """Return True if all non-empty runs in the paragraph are italic."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    return bool(runs) and all(r.italic for r in runs)


def _is_list_paragraph(paragraph) -> bool:
    style_name = getattr(paragraph.style, 'name', '') if paragraph.style else ''
    if style_name.startswith('List ') or style_name.startswith('Checklist'):
        return True
    p_pr = paragraph._p.find(qn('w:pPr'))
    if p_pr is None:
        return False
    return p_pr.find(qn('w:numPr')) is not None


def _is_candidate_after_list(text: str) -> bool:
    if not text:
        return False
    if len(text) <= 220:
        return True
    lowered = text.lower()
    return (
        lowered.startswith('then discuss:')
        or lowered.startswith('discuss:')
        or lowered.startswith('reflect:')
        or lowered.startswith('then reflect:')
        or lowered.startswith('reflection:')
        or lowered.startswith('example:')
        or lowered.startswith('examples:')
        or lowered.startswith('practice ')
        or lowered.startswith('next step:')
        or lowered.startswith('next steps:')
        or lowered.startswith('in pairs')
        or lowered.startswith('working in pairs')
    )


def _is_candidate_after_text_block(text: str) -> bool:
    if not text:
        return False
    if len(text) <= 160 and not text.endswith('.'):
        return True
    lowered = text.lower()
    return (
        lowered.startswith('practice ')
        or lowered.startswith('reflect:')
        or lowered.startswith('then reflect:')
        or lowered.startswith('reflection:')
        or lowered.startswith('example:')
        or lowered.startswith('examples:')
    )


def _find_previous_text_paragraph(paragraphs, index):
    for offset in range(index - 1, -1, -1):
        para = paragraphs[offset]
        if _normalize_text(para.text):
            return para
    return None


def _find_next_text_paragraph(paragraphs, index):
    for offset in range(index + 1, len(paragraphs)):
        para = paragraphs[offset]
        if _normalize_text(para.text):
            return para
    return None


def _apply_style_if_available(paragraph, style_name: str) -> bool:
    style = _get_style_by_name_or_id(paragraph.part.styles, style_name)
    if not style:
        return False
    paragraph.style = style
    return True


def _paragraph_style_id(paragraph) -> str:
    style = paragraph.style
    return getattr(style, 'style_id', '') if style else ''


def _paragraph_style_name(paragraph) -> str:
    style = paragraph.style
    return getattr(style, 'name', '') if style else ''


def _has_explicit_paragraph_style(paragraph) -> bool:
    p_pr = paragraph._p.find(qn('w:pPr'))
    return bool(p_pr is not None and p_pr.find(qn('w:pStyle')) is not None)


def _iter_all_paragraphs(parent) -> Iterable[Paragraph]:
    """Yield paragraphs in the document body and inside tables."""
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from _iter_all_paragraphs(cell)


def _iter_story_paragraphs(doc) -> Iterable[Paragraph]:
    yield from _iter_all_paragraphs(doc)
    for section in doc.sections:
        yield from _iter_all_paragraphs(section.header)
        yield from _iter_all_paragraphs(section.first_page_header)
        yield from _iter_all_paragraphs(section.even_page_header)
        yield from _iter_all_paragraphs(section.footer)
        yield from _iter_all_paragraphs(section.first_page_footer)
        yield from _iter_all_paragraphs(section.even_page_footer)


def _set_run_font(run, font_name: str) -> None:
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement('w:rFonts')
        r_pr.insert(0, r_fonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:eastAsia', 'w:cs'):
        r_fonts.set(qn(attr), font_name)


def _set_run_color(run, color_value: str | None) -> None:
    if not color_value:
        return
    r_pr = run._r.get_or_add_rPr()
    color = r_pr.find(qn('w:color'))
    if color is None:
        color = OxmlElement('w:color')
        r_pr.append(color)
    color.set(qn('w:val'), color_value)


def _style_color_value(style) -> str | None:
    if style is None:
        return None
    style_el = getattr(style, '_element', None)
    if style_el is None:
        return None
    color = style_el.find('.//w:rPr/w:color', NS)
    if color is None:
        color = style_el.find('.//w:color', NS)
    if color is None:
        return None
    return color.get(qn('w:val'))


def _run_color_value(run) -> str | None:
    color = run._r.find('.//w:color', NS)
    if color is None:
        return None
    return color.get(qn('w:val'))


def _insert_run_after_properties(paragraph, run) -> None:
    paragraph._p.remove(run._r)
    p_pr = paragraph._p.find(qn('w:pPr'))
    if p_pr is None:
        paragraph._p.insert(0, run._r)
    else:
        paragraph._p.insert(list(paragraph._p).index(p_pr) + 1, run._r)


def _set_run_non_italic(run) -> None:
    run.font.italic = False
    r_pr = run._r.get_or_add_rPr()
    for tag in ('w:i', 'w:iCs'):
        italic = r_pr.find(qn(tag))
        if italic is None:
            italic = OxmlElement(tag)
            r_pr.append(italic)
        italic.set(qn('w:val'), '0')


def _set_run_non_bold(run) -> None:
    run.font.bold = False
    r_pr = run._r.get_or_add_rPr()
    for tag in ('w:b', 'w:bCs'):
        bold = r_pr.find(qn(tag))
        if bold is not None:
            r_pr.remove(bold)


def _is_strong_label_run(run) -> bool:
    if run.bold is True:
        return True
    r_pr = run._r.find(qn('w:rPr'))
    if r_pr is None:
        return False
    if r_pr.find(qn('w:b')) is not None:
        return True
    r_style = r_pr.find(qn('w:rStyle'))
    if r_style is None:
        return False
    return r_style.get(qn('w:val')) in {'Strong', 'StrongChar'}


def _label_base_paragraph_style(styles):
    return _get_style_by_name_or_id(styles, 'Label Base Para') or _require_style(
        styles,
        'Label Para',
    )


def _find_semantic_label_run(paragraph):
    for index, run in enumerate(paragraph.runs):
        if not run.text or not run.text.strip():
            continue
        if _is_strong_label_run(run):
            return index, run
    return None, None


def _find_run_index(paragraph, target_run) -> int | None:
    target_el = target_run._r
    for index, run in enumerate(paragraph.runs):
        if run._r is target_el:
            return index
    return None


def _normalize_learn_label_text(style_id: str, run) -> bool:
    updated = run.text
    if style_id in LEARN_SEMANTIC_STYLE_IDS:
        match = re.match(r'^\s*Learn(?:\s*[-\u2013\u2014:]\s*|\s+)(.+?)\s*$', updated)
        if match:
            updated = match.group(1)
    updated = updated.rstrip()
    if updated.endswith(':'):
        updated = updated[:-1].rstrip()
    if updated == run.text:
        return False
    run.text = updated
    return True


def apply_semantic_div_labels(
    doc,
    reference_doc_path=None,
    tag_style: str = DIV_TAG_ICON_STYLE,
    context: PostprocessContext | None = None,
) -> int:
    """
    Add visual labels to semantic Div paragraphs emitted by the Pandoc Lua filter.

    For Div styles with a tag icon, inserts the PNG icon inline before the label text.
    Example divs (neutral/good/bad) have no icon and are handled by apply_example_block_styles.
    """
    changed = 0
    learn_base = _get_style_by_name_or_id(doc.styles, 'Learn Base')
    semantic_div_label_styles = build_semantic_div_label_styles(reference_doc_path, context)
    body = doc.element.body

    # Snapshot of direct body <w:p> children only — we will replace some with <w:tbl>.
    # Use doc._body as the parent proxy so Paragraph.style can resolve via doc.part.
    body_paras = [
        (child, Paragraph(child, doc._body))
        for child in list(body)
        if child.tag == qn('w:p')
    ]

    icon_count = 0

    for p_el, para in body_paras:
        style_id = _paragraph_style_id(para)
        has_icon = style_id in DIV_TAG_ICON_STEMS
        if not has_icon:
            continue

        if not para.runs:
            continue

        _, label_run = _find_semantic_label_run(para)
        if label_run is None:
            for idx, run in enumerate(para.runs):
                if run.text and run.text.strip():
                    label_index, label_run = idx, run
                    break
        if label_run is None:
            continue

        if _normalize_learn_label_text(style_id, label_run):
            changed += 1

        # Get color from paragraph style map (no char style needed)
        label_color = semantic_div_label_styles.get(style_id) or _run_color_value(label_run)
        # Apply color and font directly to the label run
        if label_color:
            _set_run_color(label_run, label_color)
        _set_run_font(label_run, 'Noto Sans Condensed Medium')
        _set_run_non_italic(label_run)

        if has_icon:
            icon_path = _resolve_div_tag_icon(style_id, reference_doc_path, tag_style)

            # Insert icon image + 2× NBSP inline before the label text.
            # The DivTag style (set to lowered by 4pt in the reference DOCX) aligns
            # the icon baseline with the label text without needing a 2-column table.
            if icon_path:
                tmp_para = doc.add_paragraph()
                img_run = tmp_para.add_run()
                img_run.add_picture(str(icon_path), height=DIV_TAG_ICON_HEIGHT_EMU)
                icon_r = img_run._r
                body.remove(tmp_para._p)

                # Apply DivTag character style to icon run so spacing is style-driven
                icon_rpr = icon_r.find(qn('w:rPr'))
                if icon_rpr is None:
                    icon_rpr = OxmlElement('w:rPr')
                    icon_r.insert(0, icon_rpr)
                icon_style_el = OxmlElement('w:rStyle')
                icon_style_el.set(qn('w:val'), 'DivTag')
                icon_rpr.insert(0, icon_style_el)

                # Build a run containing 2× NBSP as spacer
                nbsp_r = OxmlElement('w:r')
                nbsp_rpr = OxmlElement('w:rPr')
                nbsp_style_el = OxmlElement('w:rStyle')
                nbsp_style_el.set(qn('w:val'), 'DivTag')
                nbsp_rpr.append(nbsp_style_el)
                nbsp_r.append(nbsp_rpr)
                nbsp_t = OxmlElement('w:t')
                nbsp_t.text = '  '
                nbsp_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                nbsp_r.append(nbsp_t)

                # Prepend icon run and NBSP run before the first existing run
                first_r = p_el.find(qn('w:r'))
                if first_r is not None:
                    idx = list(p_el).index(first_r)
                    p_el.insert(idx, nbsp_r)
                    p_el.insert(idx, icon_r)
                else:
                    p_el.append(icon_r)
                    p_el.append(nbsp_r)

            icon_count += 1
            changed += 1

        if style_id in LEARN_SEMANTIC_STYLE_IDS:
            if learn_base is not None and getattr(para.style, 'name', None) != learn_base.name:
                para.style = learn_base
                changed += 1

    if icon_count:
        print(f'  apply_semantic_div_labels: {icon_count} icon labels')
    return changed


def apply_body_text_to_normal_paragraphs(doc) -> int:
    """Make ordinary prose explicit Body Text instead of implicit/explicit Normal."""
    body_text = _require_style(doc.styles, 'Body Text')
    body_text_style_id = getattr(body_text, 'style_id', 'BodyText')

    changed = 0
    for p_el in doc.element.body.iterchildren():
        if p_el.tag != qn('w:p'):
            continue

        text = _normalize_text(''.join(t.text or '' for t in p_el.iter(qn('w:t'))))
        if not text:
            continue

        p_pr = p_el.find(qn('w:pPr'))
        p_style = p_pr.find(qn('w:pStyle')) if p_pr is not None else None
        style_id = p_style.get(qn('w:val')) if p_style is not None else ''
        num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None

        if num_pr is not None:
            continue
        if style_id == body_text_style_id:
            continue
        if style_id.startswith(('Heading', 'DivLabel', 'List', 'Checklist')):
            continue
        if style_id and style_id != 'Normal':
            continue

        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            p_el.insert(0, p_pr)
        if p_style is None:
            p_style = OxmlElement('w:pStyle')
            p_pr.insert(0, p_style)
        p_style.set(qn('w:val'), body_text_style_id)
        changed += 1
    return changed


def strip_activity_codes_from_headings(doc) -> int:
    """Remove activity-code suffixes from visible heading/title rows."""
    changed = 0
    for para in doc.paragraphs:
        if not _is_heading(para):
            continue
        text = _normalize_text(para.text)
        updated = ACTIVITY_CODE_SUFFIX_RE.sub('', text).strip()
        if updated == text:
            continue
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = updated
        else:
            para.add_run(updated)
        changed += 1
    return changed


def remove_pandoc_generated_styles(doc) -> int:
    """Replace Pandoc fallback styles that are not part of the reference DOCX."""
    body_text = _require_style(doc.styles, 'Body Text')
    body_text_char = _require_style(doc.styles, 'Body Text Char')
    changed = 0

    for para in _iter_all_paragraphs(doc):
        p_pr = para._p.find(qn('w:pPr'))
        p_style = p_pr.find(qn('w:pStyle')) if p_pr is not None else None
        if p_style is not None and p_style.get(qn('w:val')) == 'Compact':
            para.style = body_text
            changed += 1

        for run in para.runs:
            r_pr = run._r.find(qn('w:rPr'))
            r_style = r_pr.find(qn('w:rStyle')) if r_pr is not None else None
            if r_style is not None and r_style.get(qn('w:val')) == 'VerbatimChar':
                run.style = body_text_char
                changed += 1

    return changed


def _style_ids(styles) -> set[str]:
    return {
        getattr(style, 'style_id', '')
        for style in styles
        if getattr(style, 'style_id', '')
    }


def purge_styles_not_in_reference(
    doc,
    reference_doc_path,
    context: PostprocessContext | None = None,
) -> int:
    """
    Remove generated DOCX style references and style definitions that are not
    present in the reference DOCX.
    """
    if not reference_doc_path:
        return 0
    allowed_style_ids = context.reference_style_ids() if context is not None else _style_ids(Document(reference_doc_path).styles)
    if not allowed_style_ids:
        return 0
    body_text = _require_style(doc.styles, 'Body Text')
    changed = 0

    for para in _iter_story_paragraphs(doc):
        para_style_id = _paragraph_style_id(para)
        if para_style_id and para_style_id not in allowed_style_ids:
            para.style = body_text
            changed += 1

        for run in para.runs:
            r_pr = run._r.find(qn('w:rPr'))
            r_style = r_pr.find(qn('w:rStyle')) if r_pr is not None else None
            style_id = r_style.get(qn('w:val')) if r_style is not None else ''
            if style_id and style_id not in allowed_style_ids:
                r_pr.remove(r_style)
                changed += 1

    for style in list(doc.styles):
        style_id = getattr(style, 'style_id', '')
        if style_id and style_id not in allowed_style_ids:
            style.delete()
            changed += 1

    return changed


def disable_heading_style_page_breaks(doc) -> int:
    """Let section breaks, not heading styles, control module/unit page starts."""
    changed = 0
    for style_name in ('Heading 1', 'Heading 2'):
        style = _get_style_by_name_or_id(doc.styles, style_name)
        if style is None:
            continue
        style_el = getattr(style, '_element', None)
        if style_el is None:
            continue
        p_pr = style_el.find(qn('w:pPr'))
        page_break = p_pr.find(qn('w:pageBreakBefore')) if p_pr is not None else None
        if p_pr is not None and page_break is not None:
            p_pr.remove(page_break)
            changed += 1
    return changed


def insert_section_breaks_before_modules_and_units(doc) -> int:
    """Start every module and unit in its own Word section."""
    boundaries = []
    for para in doc.paragraphs:
        if _paragraph_style_name(para) not in {'Heading 1', 'Heading 2'}:
            continue
        text = _normalize_text(para.text)
        if not (
            MODULE_HEADING_RE.match(text)
            or UNIT_HEADING_RE.match(text)
            or MODULE_REVIEW_HEADING_RE.match(text)
        ):
            continue
        boundaries.append(para)

    changed = 0
    for para in reversed(boundaries[1:]):
        if _insert_section_break_before_paragraph(para):
            changed += 1
    return changed


def apply_page_breaks_before_modules_and_units(doc) -> int:
    """Backward-compatible wrapper for the old page-start function name."""
    return insert_section_breaks_before_modules_and_units(doc)


def normalize_module_headings(doc) -> int:
    """Strip unit ranges from module titles while preserving Heading 1 level."""
    heading_1 = _require_style(doc.styles, 'Heading 1')
    changed = 0
    for para in doc.paragraphs:
        text = _normalize_text(para.text)
        if not MODULE_HEADING_RE.match(text):
            continue
        if _paragraph_style_name(para) != 'Heading 1':
            para.style = heading_1
            changed += 1
        cleaned = MODULE_UNIT_RANGE_SUFFIX_RE.sub('', text)
        if cleaned != text and _replace_paragraph_text_preserve_format(para, cleaned):
            changed += 1
    return changed


def _find_unit_tile_prototype(reference_doc_path, context: PostprocessContext | None = None) -> object | None:
    if context is not None:
        return context.unit_tile_prototype()
    if not reference_doc_path:
        return None
    ref_path = Path(reference_doc_path)
    if not ref_path.exists():
        return None
    try:
        ref_doc = Document(ref_path)
    except Exception:
        return None
    for table in ref_doc.tables:
        try:
            first = _clean_word_text(table.cell(0, 0).text)
            second = _clean_word_text(table.cell(0, 1).text)
        except Exception:
            continue
        if re.match(r'^(?:U0|U#|#|\d+)$', first) and 'Unit Title' in second:
            return copy.deepcopy(table._tbl)
    return None


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        paragraph._p.remove(child)


def _make_paragraph_normal(paragraph) -> None:
    try:
        paragraph.style = paragraph.part.styles['Normal']
    except Exception:
        pass


def _make_page_break_paragraph_like(paragraph) -> None:
    _make_paragraph_normal(paragraph)
    _clear_paragraph(paragraph)
    _make_paragraph_normal(paragraph)
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        p_pr.append(spacing)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    if p_pr.find(qn('w:pageBreakBefore')) is None:
        p_pr.append(OxmlElement('w:pageBreakBefore'))


def _hide_heading_paragraph_keep_for_styleref(paragraph, keep_page_break: bool = True) -> None:
    """Hide a heading visually while keeping its text available to STYLEREF."""
    p_pr = paragraph._p.get_or_add_pPr()
    spacing = p_pr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        p_pr.append(spacing)
    spacing.set(qn('w:before'), '0')
    spacing.set(qn('w:after'), '0')
    if keep_page_break and p_pr.find(qn('w:pageBreakBefore')) is None:
        p_pr.append(OxmlElement('w:pageBreakBefore'))

    r_pr = p_pr.find(qn('w:rPr'))
    if r_pr is None:
        r_pr = OxmlElement('w:rPr')
        p_pr.append(r_pr)
    for tag in ('w:vanish', 'w:specVanish'):
        if r_pr.find(qn(tag)) is None:
            r_pr.append(OxmlElement(tag))

    for run in paragraph.runs:
        run_pr = run._r.get_or_add_rPr()
        for tag in ('w:vanish', 'w:specVanish'):
            if run_pr.find(qn(tag)) is None:
                run_pr.append(OxmlElement(tag))


def _new_unit_tile_table(doc, prototype_tbl):
    if prototype_tbl is None:
        raise RuntimeError(
            'Unit title table prototype was not found in the reference DOCX.'
        )
    return copy.deepcopy(prototype_tbl)


def _replace_cell_text_preserve_format(cell, text: str) -> None:
    """
    Replace visible text without rebuilding the cell content, preserving the
    cloned prototype's table, cell, paragraph, and run formatting.
    """
    text_nodes = cell._tc.findall('.//w:t', NS)
    if not text_nodes:
        cell.paragraphs[0].add_run(text)
        return
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ''


def _replace_paragraph_text_preserve_format(paragraph, text: str) -> bool:
    text_nodes = paragraph._p.findall('.//w:t', NS)
    if not text_nodes:
        paragraph.add_run(text)
        return True
    old_text = ''.join(node.text or '' for node in text_nodes)
    if old_text == text:
        return False
    text_nodes[0].text = text
    for node in text_nodes[1:]:
        node.text = ''
    return True


def _set_header_text(header, text: str) -> bool:
    if not text:
        return False
    changed = False
    target = None
    for paragraph in header.paragraphs:
        paragraph_text = _normalize_text(paragraph.text)
        if paragraph_text in {'Module X Title', 'Unit Y Title'}:
            target = paragraph
            break
        if target is None and paragraph_text:
            target = paragraph
    if target is None:
        target = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    if _replace_paragraph_text_preserve_format(target, text):
        changed = True
    return changed


def _set_header_styleref(header, style_name: str = 'Heading 2') -> bool:
    target = None
    for paragraph in header.paragraphs:
        if _normalize_text(paragraph.text):
            target = paragraph
            break
    if target is None:
        target = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

    p_pr = target._p.find(qn('w:pPr'))
    saved_p_pr = copy.deepcopy(p_pr) if p_pr is not None else None
    for child in list(target._p):
        target._p.remove(child)
    if saved_p_pr is not None:
        target._p.append(saved_p_pr)

    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), f'STYLEREF  "{style_name}"  \\* MERGEFORMAT')
    run = OxmlElement('w:r')
    text = OxmlElement('w:t')
    text.text = style_name
    run.append(text)
    fld.append(run)
    target._p.append(fld)
    return True


def _module_unit_contexts(doc) -> list[tuple[str, str]]:
    contexts: list[tuple[str, str]] = []
    current_module = ''
    current_unit = ''
    for para in doc.paragraphs:
        style_name = _paragraph_style_name(para)
        if style_name not in {'Heading 1', 'Heading 2'}:
            continue
        text = _normalize_text(para.text)
        if style_name == 'Heading 1' and MODULE_HEADING_RE.match(text):
            current_module = text
            current_unit = ''
            contexts.append((current_module, current_unit))
        elif style_name == 'Heading 2' and UNIT_HEADING_RE.match(text):
            current_unit = text
            contexts.append((current_module, current_unit))
        elif style_name == 'Heading 2' and MODULE_REVIEW_HEADING_RE.match(text):
            contexts.append((current_module, text))
    return contexts


def update_running_headers(doc) -> int:
    """Write module/unit running headers directly into each Word section."""
    contexts = _module_unit_contexts(doc)
    if not contexts:
        return 0

    changed = 0
    for idx, section in enumerate(doc.sections):
        module_title, unit_title = contexts[min(idx, len(contexts) - 1)]
        odd_header = module_title or unit_title
        even_header = unit_title or module_title

        section.header.is_linked_to_previous = False
        section.even_page_header.is_linked_to_previous = False
        section.first_page_header.is_linked_to_previous = False
        if _set_header_text(section.header, odd_header):
            changed += 1
        if _set_header_text(section.first_page_header, odd_header):
            changed += 1
        if _set_header_text(section.even_page_header, even_header):
            changed += 1
    return changed


def replace_unit_headings_with_title_tables(
    doc,
    reference_doc_path=None,
    context: PostprocessContext | None = None,
) -> int:
    """
    Replace `Unit N - Title` Heading 2 paragraphs with the reference unit-title
    table prototype. This avoids slow Word COM Quick Part insertion.
    """
    prototype_tbl = _find_unit_tile_prototype(reference_doc_path, context)
    if prototype_tbl is None:
        return 0
    changed = 0
    for para in list(doc.paragraphs):
        match = UNIT_HEADING_RE.match(_normalize_text(para.text))
        if not match:
            continue

        tbl = _new_unit_tile_table(doc, prototype_tbl)
        para._p.addnext(tbl)
        table = Table(tbl, para._parent)
        _replace_cell_text_preserve_format(table.cell(0, 0), str(int(match.group(1))))
        _replace_cell_text_preserve_format(table.cell(0, 1), match.group(2))
        para._element.getparent().remove(para._element)
        changed += 1
    return changed


def _iter_body_blocks(doc):
    for child in doc.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def _is_unit_title_table(table) -> bool:
    try:
        first_row = table.rows[0]
        first_cell = _clean_word_text(first_row.cells[0].text)
    except Exception:
        return False
    return bool(re.match(r'^(?:U)?\d+$', first_cell))


def restore_unit_overview_headings_after_unit_tiles(doc) -> int:
    """
    After unit title Quick Parts are inserted, ensure the following
    `Unit Overview` paragraph keeps its heading style.
    """
    changed = 0
    blocks = list(_iter_body_blocks(doc))
    for index, block in enumerate(blocks[:-1]):
        if not isinstance(block, Table) or not _is_unit_title_table(block):
            continue
        next_block = blocks[index + 1]
        if not isinstance(next_block, Paragraph):
            continue
        if _normalize_text(next_block.text) != 'Unit Overview':
            continue
        if _apply_style_if_available(next_block, 'Heading 3'):
            changed += 1
    return changed


def _can_apply_after_list_to_paragraph(paragraph) -> bool:
    """
    Only apply After List to paragraphs that are still body-like.

    This must not override heading styles or other already-semantic paragraph
    styles coming from Pandoc/reference.docx.
    """
    if _is_heading(paragraph):
        return False
    style_name = getattr(paragraph.style, 'name', '') if paragraph.style else ''
    return style_name in {'Normal', 'Body Text', 'Block Text', 'After List'}


def _is_short_follow_on_context(paragraph) -> bool:
    if _is_heading(paragraph) or _is_list_paragraph(paragraph):
        return False
    style_name = getattr(paragraph.style, 'name', '') if paragraph.style else ''
    return style_name in {'Normal', 'Body Text', 'Block Text', 'After List'}


def _is_module_homework_target(paragraphs, index, text: str) -> bool:
    lowered = text.lower()
    if 'homework target:' not in lowered:
        return False
    if not WORD_COUNT_RE.search(text):
        return False
    prev_para = _find_previous_text_paragraph(paragraphs, index)
    if prev_para is None or not _is_heading(prev_para):
        return False
    style_name = getattr(prev_para.style, 'name', '') if prev_para.style else ''
    return style_name == 'Heading 2'


def _is_neutral_model_example(para, text: str) -> bool:
    """Neutral writing examples arrive from Pandoc as built-in quote styles."""
    if _is_heading(para) or _is_list_paragraph(para):
        return False
    style_name = _paragraph_style_name(para)
    if style_name in NEUTRAL_MODEL_SOURCE_STYLES:
        return True
    if style_name in {'Model', 'Model Bad', 'Model Good'}:
        return False
    return bool(QUOTED_MODEL_RE.match(text))


def apply_semantic_styles(doc):
    """
    Apply semantic paragraph styles that do not require Word COM.

    Quick Parts are handled separately in a Word automation pass so they can be
    inserted as real Building Blocks from the template.
    """
    changed = 0
    model_mode = None
    prev_text_para = None

    for para in doc.paragraphs:
        text = _normalize_text(para.text)
        if not text:
            continue

        prev_para = prev_text_para
        prev_text_para = para

        if text in BAD_MODEL_HEADINGS:
            model_mode = 'bad'
            continue

        if text in GOOD_MODEL_HEADINGS:
            model_mode = 'good'
            continue

        if _is_heading(para):
            model_mode = None
            continue

        if model_mode == 'bad' and _apply_style_if_available(para, 'Model Bad'):
            changed += 1
            model_mode = None
            continue

        if model_mode == 'good' and _apply_style_if_available(para, 'Model Good'):
            changed += 1
            model_mode = None
            continue

        if _is_neutral_model_example(para, text) and _apply_style_if_available(para, 'Model'):
            changed += 1
            continue

        is_homework_target = False
        if prev_para is not None and WORD_COUNT_RE.search(text) and 'homework target:' in text.lower():
            prev_style_name = getattr(prev_para.style, 'name', '') if prev_para.style else ''
            is_homework_target = _is_heading(prev_para) and prev_style_name == 'Heading 2'

        if is_homework_target and _apply_style_if_available(para, 'Homework Target'):
            changed += 1
            continue

        if prev_para is not None:
            should_apply_after_list = False
            if _is_list_paragraph(prev_para) and _is_candidate_after_list(text):
                should_apply_after_list = True
            elif _is_short_follow_on_context(prev_para) and _is_candidate_after_text_block(text):
                should_apply_after_list = True

            if (
                should_apply_after_list
                and _can_apply_after_list_to_paragraph(para)
                and _apply_style_if_available(para, 'After List')
            ):
                changed += 1

    return changed


def apply_list_styles(
    doc,
    bullet_style='List Bullet 2',
    number_style='List Number 2',
    alpha_style='List Number 3',
):
    """Apply specific styles to bullet and numbered list paragraphs."""
    num_map = _build_num_format_map(doc)
    styles = doc.styles
    bullet = _get_style_by_name_or_id(styles, bullet_style) or _get_style_by_name_or_id(
        styles,
        bullet_style.replace(' ', ''),
    )
    number = _get_style_by_name_or_id(styles, number_style) or _get_style_by_name_or_id(
        styles,
        number_style.replace(' ', ''),
    )
    alpha = _get_style_by_name_or_id(styles, alpha_style) or _get_style_by_name_or_id(
        styles,
        alpha_style.replace(' ', ''),
    )
    if not bullet:
        raise RuntimeError(f'Required style is missing from the reference DOCX: {bullet_style}')
    if not number:
        raise RuntimeError(f'Required style is missing from the reference DOCX: {number_style}')
    if not alpha:
        raise RuntimeError(f'Required style is missing from the reference DOCX: {alpha_style}')

    applied = 0
    alpha_base_num_id = _num_id_from_style(alpha)
    current_alpha_num_id = None
    expected_alpha_value = None
    for para in doc.paragraphs:
        normalized_text = _normalize_text(para.text)
        alpha_marker = ALPHA_ORDINAL_RE.match(normalized_text)
        p_pr = para._p.find(qn('w:pPr'))
        num_pr = p_pr.find(qn('w:numPr')) if p_pr is not None else None
        if num_pr is not None:
            num_id_el = num_pr.find(qn('w:numId'))
            ilvl_el = num_pr.find(qn('w:ilvl'))
            if num_id_el is None:
                continue
            num_id = num_id_el.get(qn('w:val'))
            ilvl = ilvl_el.get(qn('w:val')) if ilvl_el is not None else '0'
            fmt = None
            if num_id in num_map:
                fmt = num_map[num_id].get(ilvl)
            if fmt == 'bullet' and bullet:
                para.style = bullet
                applied += 1
                continue
            if fmt in {'upperLetter', 'lowerLetter', 'upperAlpha', 'lowerAlpha'}:
                para.style = alpha
                applied += 1
                current_alpha_num_id = num_id
                expected_alpha_value = None
                continue
            if fmt and fmt != 'bullet' and number:
                para.style = number
                applied += 1
                current_alpha_num_id = None
                expected_alpha_value = None
                continue

        if alpha and alpha_marker:
            style_name = getattr(para.style, 'name', '') if para.style else ''
            if not style_name.startswith('Heading'):
                para.style = alpha
                applied += 1
                marker_value = _alpha_marker_start_value(alpha_marker.group(1))
                if (
                    alpha_base_num_id
                    and (current_alpha_num_id is None or expected_alpha_value != marker_value)
                ):
                    current_alpha_num_id = _create_restarted_num_id(
                        doc,
                        alpha_base_num_id,
                        marker_value,
                    )
                if current_alpha_num_id:
                    _apply_direct_num_id(para, current_alpha_num_id)
                expected_alpha_value = marker_value + 1
                continue

        if normalized_text:
            current_alpha_num_id = None
            expected_alpha_value = None
    return applied


def _remove_prefix_from_runs(paragraph, prefix_len: int) -> bool:
    remaining = prefix_len
    changed = False
    for run in paragraph.runs:
        if remaining <= 0:
            break
        text = run.text
        if not text:
            continue
        if len(text) <= remaining:
            run.text = ''
            remaining -= len(text)
            changed = True
            continue
        run.text = text[remaining:]
        remaining = 0
        changed = True
        break
    return changed and remaining == 0


def strip_literal_alpha_markers(doc, alpha_style='List Number 3') -> int:
    """Remove source-text A./B. markers after alphabetic list styling is applied."""
    alpha = _get_style_by_name_or_id(doc.styles, alpha_style) or _get_style_by_name_or_id(
        doc.styles,
        alpha_style.replace(' ', ''),
    )
    if not alpha:
        raise RuntimeError(f'Required style is missing from the reference DOCX: {alpha_style}')

    changed = 0
    alpha_style_id = getattr(alpha, 'style_id', None)
    alpha_style_name = getattr(alpha, 'name', None)
    marker_re = re.compile(r'^\s*[A-Z][\.)]\s+')
    for para in doc.paragraphs:
        para_style = para.style
        if para_style is None:
            continue
        if (
            getattr(para_style, 'style_id', None) != alpha_style_id
            and getattr(para_style, 'name', None) != alpha_style_name
        ):
            continue
        match = marker_re.match(para.text)
        if not match:
            continue
        if _remove_prefix_from_runs(para, len(match.group(0))):
            changed += 1
    return changed


def _list_winword_pids() -> set[int]:
    try:
        result = subprocess.run(
            ['tasklist', '/FI', 'IMAGENAME eq WINWORD.EXE', '/FO', 'CSV', '/NH'],
            capture_output=True,
            text=True,
            check=False,
        )
    except OSError:
        return set()

    pids: set[int] = set()
    reader = csv.reader(line for line in result.stdout.splitlines() if line.strip())
    for row in reader:
        if len(row) < 2 or row[0].strip('"').upper() == 'INFO:':
            continue
        try:
            pids.add(int(row[1]))
        except ValueError:
            continue
    return pids


def _get_process_id_from_window(hwnd: int) -> int | None:
    if not hwnd:
        return None
    pid = ctypes.c_ulong()
    ctypes.windll.user32.GetWindowThreadProcessId(hwnd, ctypes.byref(pid))
    return int(pid.value or 0) or None


def _get_building_block_entry(word_app, template_path: Path, entry_name: str):
    template = word_app.Templates(str(template_path))
    return template.BuildingBlockEntries(entry_name)


def _get_first_available_building_block_entry(word_app, template_path: Path, entry_names) -> object:
    names = entry_names if isinstance(entry_names, (tuple, list)) else (entry_names,)
    last_exc = None
    for name in names:
        try:
            return _get_building_block_entry(word_app, template_path, name)
        except Exception as exc:
            last_exc = exc
    if last_exc is not None:
        raise last_exc
    raise RuntimeError('No Building Block names were provided.')


def _replace_cell_text(cell, text: str) -> None:
    cell.Range.Text = text


def _insert_building_block_after_paragraph(entry, paragraph, add_page_break=False):
    target_range = paragraph.Range.Duplicate
    target_range.Collapse(WORD_COLLAPSE_END)
    target_range.InsertParagraphAfter()
    try:
        target_range.Style = paragraph.Application.ActiveDocument.Styles('Normal')
    except Exception:
        pass
    if add_page_break:
        try:
            target_range.InsertBreak(WORD_PAGE_BREAK)
            target_range.Style = paragraph.Application.ActiveDocument.Styles('Normal')
        except Exception:
            pass
    target_range.Collapse(WORD_COLLAPSE_END)
    before_start = target_range.Start
    entry.Insert(Where=target_range, RichText=True)
    return before_start


def _find_table_inserted_after_position(doc, position: int):
    for index in range(1, doc.Tables.Count + 1):
        table = doc.Tables(index)
        if table.Range.Start >= position:
            return table
    return None


def _insert_unit_tile_building_block(paragraph, template_path: Path) -> bool:
    match = UNIT_HEADING_RE.match(_clean_word_text(paragraph.Range.Text))
    if not match:
        return False
    paragraph_range = paragraph.Range.Duplicate
    entry = _get_first_available_building_block_entry(
        paragraph.Application,
        template_path,
        BUILDING_BLOCK_NAMES['unit'],
    )
    insert_start = _insert_building_block_after_paragraph(entry, paragraph, add_page_break=True)
    table = _find_table_inserted_after_position(paragraph.Application.ActiveDocument, insert_start)
    if table is None:
        return False
    _replace_cell_text(table.Cell(1, 1), str(int(match.group(1))))
    _replace_cell_text(table.Cell(1, 2), match.group(2))
    paragraph_range.Delete()
    return True


def _can_apply_quick_parts(template_path: Path | None) -> bool:
    return bool(template_path and template_path.exists() and pythoncom is not None and win32com is not None)


def apply_quick_parts(docx_path, template_path: Path | None = None) -> int:
    """
    Insert real Quick Parts using a dedicated Word COM instance.

    Safety rules:
    - check for existing WINWORD.exe processes before automation starts
    - create a fresh Word instance via DispatchEx instead of attaching to an
      existing interactive session
    - close only the document and Word instance created by this function
    - never kill or terminate unrelated WINWORD.exe processes
    """
    if not _can_apply_quick_parts(template_path):
        return 0

    preexisting_pids = _list_winword_pids()
    assert pythoncom is not None
    assert win32com is not None
    pythoncom.CoInitialize()  # pylint: disable=no-member
    word = None
    doc = None
    owned_pid = None
    changes = 0
    try:
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = WORD_ALERTS_NONE
        owned_pid = _get_process_id_from_window(getattr(word, 'Hwnd', 0))
        if owned_pid is not None and owned_pid in preexisting_pids:
            raise RuntimeError(
                'Word automation attached to an existing WINWORD.exe process; aborting Quick Part insertion.'
            )

        doc = word.Documents.Open(str(Path(docx_path).resolve()), AddToRecentFiles=False)
        para_index = doc.Paragraphs.Count
        while para_index >= 1:
            paragraph = doc.Paragraphs(para_index)
            text = _clean_word_text(paragraph.Range.Text)
            if not text:
                para_index -= 1
                continue

            if UNIT_HEADING_RE.match(text):
                if _insert_unit_tile_building_block(paragraph, template_path):
                    changes += 1
                para_index -= 1
                continue

            para_index -= 1

        if changes > 0:
            doc.Save()
        return changes
    finally:
        if doc is not None:
            doc.Close(SaveChanges=WORD_DO_NOT_SAVE_CHANGES)
        if word is not None:
            word.Quit(SaveChanges=WORD_DO_NOT_SAVE_CHANGES)
        if pythoncom is not None:
            getattr(pythoncom, 'CoUninitialize', lambda: None)()


def apply_checklist_style(doc) -> int:
    """Apply Checklist style to list bullet items inside any edit div.

    When Pandoc converts '- [ ]' task list syntax with a reference DOCX that defines
    List Bullet 2, the output is identical to plain '- ' bullets (List Bullet 2 style,
    no checkbox glyph). We therefore apply Checklist style to all List Bullet 2 / List
    Bullet paragraphs that follow any DivLabelEdit paragraph.
    """
    checklist_style = _get_style_by_name_or_id(doc.styles, 'Checklist')
    if not checklist_style:
        return 0

    paragraphs = list(doc.paragraphs)
    changed = 0
    in_edit = False

    for para in paragraphs:
        style_id = _paragraph_style_id(para)

        if style_id == 'DivLabelEdit':
            in_edit = True
            continue

        if in_edit:
            style_name = getattr(para.style, 'name', '') if para.style else ''
            if style_name in ('List Bullet 2', 'List Bullet', 'List Bullet 3'):
                para.style = checklist_style
                changed += 1
            elif _is_heading(para) or (style_id and style_id.startswith('DivLabel')):
                in_edit = False

    return changed


def apply_example_block_styles(doc) -> int:
    """Apply AW Example Good / AW Example Bad / AW Example styles after div label paragraphs.

    Styles every paragraph (quoted text, italic prose, lists) from a DivLabelExample*
    label up to but not including the first non-italic Body Text / Normal paragraph,
    which is treated as the task instruction following the example. Stops immediately
    at the next DivLabel* or heading.

    Rule:
    - italic Body Text / Normal → example content, continue
    - list paragraph → example content, continue (unless closing quote already seen)
    - Block Text / quoted styles → example content; if text ends with closing quote,
      set closing-quote-seen flag so the NEXT non-empty paragraph stops styling
    - non-italic Body Text / Normal with text → task instruction boundary; stop
    - empty paragraph → skip (do not style, do not stop)
    - heading or DivLabel* → stop

    Closing-quote rule: a paragraph whose text ends with " or " signals the end of
    the quoted example body. Anything after it (lists, task instructions) is not part
    of the example and should not receive example formatting.

    Hidden-label rule: if the DivLabelExample* paragraph text is exactly `No Title`,
    treat it as a control marker only. Remove that label paragraph from the document,
    but keep the neutral/good/bad example style active for the following content.
    """
    good_style = _get_style_by_name_or_id(doc.styles, 'AW Example Good')
    bad_style = _get_style_by_name_or_id(doc.styles, 'AW Example Bad')
    neutral_style = _get_style_by_name_or_id(doc.styles, 'AW Example')
    if not good_style and not bad_style:
        return 0

    paragraphs = list(doc.paragraphs)
    changed = 0
    target_style = None
    _after_closing_quote = False
    # Track whether italic content has been seen in the current example div.
    # ADV-style examples use italic for example text; stopping at the first
    # non-italic paragraph correctly separates example from task instruction.
    # INT-style examples use plain (non-italic) text throughout; in that case
    # _seen_italic stays False and non-italic paragraphs are styled until the
    # next DivLabel or heading.
    _seen_italic = False

    for para in paragraphs:
        style_id = _paragraph_style_id(para)
        text = _normalize_text(para.text)

        if style_id == 'DivLabelExampleGood':
            target_style = good_style
            _after_closing_quote = False
            _seen_italic = False
            if NO_TITLE_MARKER_RE.match(text):
                para._element.getparent().remove(para._element)
                changed += 1
            continue
        if style_id == 'DivLabelExampleBad':
            target_style = bad_style
            _after_closing_quote = False
            _seen_italic = False
            if NO_TITLE_MARKER_RE.match(text):
                para._element.getparent().remove(para._element)
                changed += 1
            continue
        if style_id == 'DivLabelExample':
            target_style = neutral_style
            _after_closing_quote = False
            _seen_italic = False
            if NO_TITLE_MARKER_RE.match(text):
                para._element.getparent().remove(para._element)
                changed += 1
            continue

        if target_style is None:
            continue

        if _is_heading(para) or (style_id and style_id.startswith('DivLabel')):
            target_style = None
            _after_closing_quote = False
            _seen_italic = False
            continue

        style_name = getattr(para.style, 'name', '') if para.style else ''

        if not text:
            # Empty paragraph — skip without stopping
            continue

        # After a closing quote, stop styling — lists and task instructions that
        # follow the quoted block are not part of the example.
        if _after_closing_quote:
            target_style = None
            _after_closing_quote = False
            _seen_italic = False
            continue

        if _is_list_paragraph(para):
            para.style = target_style
            changed += 1
        elif style_name in NEUTRAL_MODEL_SOURCE_STYLES or bool(QUOTED_MODEL_RE.match(text)):
            para.style = target_style
            changed += 1
            # If this paragraph ends with a closing quote, flag the boundary
            if text.endswith(('”', '”')):
                _after_closing_quote = True
        elif style_name in ('Body Text', 'Normal') and _is_paragraph_italic(para):
            para.style = target_style
            changed += 1
            _seen_italic = True
            if text.endswith(('”', '”')):
                _after_closing_quote = True
        elif style_name in ('Body Text', 'Normal') and not _is_paragraph_italic(para):
            if _seen_italic:
                # ADV-style: italic content seen earlier — this non-italic paragraph
                # is the task instruction boundary; stop without styling it.
                target_style = None
                _after_closing_quote = False
                _seen_italic = False
            else:
                # INT-style: no italic seen yet — plain-text example content;
                # style it and keep going until the next DivLabel or heading.
                para.style = target_style
                changed += 1

    return changed


EXAMPLE_CONTENT_STYLE_NAMES = frozenset({
    'AW Example',
    'AW Example Good',
    'AW Example Bad',
    # Pandoc/reference DOCX block quotes can remain standard Block Text when
    # the source has not mapped them to an AW Example style. Treat them as the
    # same kind of boxed/quoted content for follow-on spacing only.
    'Block Text',
})

EXAMPLE_CONTENT_STYLE_IDS = frozenset({
    'AWExample',
    'AWExampleGood',
    'AWExampleBad',
    'BlockText',
})

BODY_LIKE_STYLE_NAMES = frozenset({
    'Normal',
    'Body Text',
    'After List',
})


def _is_example_content_paragraph(paragraph) -> bool:
    style_name = _paragraph_style_name(paragraph)
    style_id = _paragraph_style_id(paragraph)
    return style_name in EXAMPLE_CONTENT_STYLE_NAMES or style_id in EXAMPLE_CONTENT_STYLE_IDS


def _spacing_before_twips_from_ppr(p_pr) -> int | None:
    if p_pr is None:
        return None
    spacing = p_pr.find(qn('w:spacing'))
    if spacing is None:
        return None
    before = spacing.get(qn('w:before'))
    if before is None:
        return None
    try:
        return int(before)
    except ValueError:
        return None


def _style_spacing_before_twips(style) -> int | None:
    """Return the nearest explicit style-level space-before value in twips."""
    seen = set()
    current = style
    while current is not None:
        style_id = getattr(current, 'style_id', None)
        if style_id in seen:
            break
        if style_id is not None:
            seen.add(style_id)
        p_pr = current.element.find(qn('w:pPr')) if current.element is not None else None
        value = _spacing_before_twips_from_ppr(p_pr)
        if value is not None:
            return value
        current = getattr(current, 'base_style', None)
    return None


def _paragraph_direct_spacing_before_twips(paragraph) -> int | None:
    p_pr = paragraph._p.find(qn('w:pPr'))
    return _spacing_before_twips_from_ppr(p_pr)


def _paragraph_effective_spacing_before_twips(paragraph) -> int | None:
    direct = _paragraph_direct_spacing_before_twips(paragraph)
    if direct is not None:
        return direct
    style = paragraph.style
    if style is None:
        return None
    return _style_spacing_before_twips(style)


def _set_paragraph_spacing_before_twips(paragraph, before_twips: int) -> None:
    p_pr = paragraph._p.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        paragraph._p.insert(0, p_pr)
    spacing = p_pr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        p_pr.append(spacing)
    spacing.set(qn('w:before'), str(before_twips))


def _is_example_following_spacing_candidate(paragraph) -> bool:
    if _is_heading(paragraph) or _is_list_paragraph(paragraph):
        return False
    style_id = _paragraph_style_id(paragraph)
    if style_id and style_id.startswith('DivLabel'):
        return False
    if _is_example_content_paragraph(paragraph):
        return False
    style_name = _paragraph_style_name(paragraph)
    if style_name in BODY_LIKE_STYLE_NAMES:
        return True
    # Some source titles, prompts, or fallback styles are still ordinary prose.
    # They can receive spacing if their own style does not already specify a
    # larger gap; semantic labels and headings are excluded above.
    return not style_name.startswith(('Heading', 'List ', 'Checklist'))


def apply_spacing_after_example_blocks(doc, space_before_twips: int = 120) -> int:
    """Add a small gap before body prose that follows an example block.

    The example content itself may render as bordered paragraphs or table-like
    boxes, so paragraph space-after on AW Example styles is not reliable. Apply
    direct space-before to the first following prose paragraph instead, while
    preserving headings and any style/direct formatting that already creates a
    larger gap.
    """
    paragraphs = list(doc.paragraphs)
    changed = 0
    after_example = False

    for para in paragraphs:
        text = _normalize_text(para.text)
        if not text:
            continue

        if _is_example_content_paragraph(para):
            after_example = True
            continue

        if not after_example:
            continue

        after_example = False
        if not _is_example_following_spacing_candidate(para):
            continue

        existing = _paragraph_effective_spacing_before_twips(para)
        if existing is not None and existing >= space_before_twips:
            continue

        _set_paragraph_spacing_before_twips(para, space_before_twips)
        changed += 1

    return changed


def apply_spacing_after_lists(doc, space_after_twips: int = 280) -> int:
    """Add space-after to the last list paragraph in each list run.

    Sets w:spacing/@w:after on the last list item directly rather than modifying
    the following prose paragraph, which avoids overriding existing space-before
    values on non-list styles.
    """
    paragraphs = list(doc.paragraphs)
    changed = 0
    next_text_para = None

    for para in reversed(paragraphs):
        text = _normalize_text(para.text)
        if not _is_list_paragraph(para):
            if text:
                next_text_para = para
            continue

        if next_text_para is not None and _is_list_paragraph(next_text_para):
            if text:
                next_text_para = para
            continue  # not the last item in the run

        p_pr = para._p.find(qn('w:pPr'))
        if p_pr is None:
            p_pr = OxmlElement('w:pPr')
            para._p.insert(0, p_pr)
        spacing = p_pr.find(qn('w:spacing'))
        if spacing is None:
            spacing = OxmlElement('w:spacing')
            p_pr.append(spacing)
        current = spacing.get(qn('w:after'))
        if current != str(space_after_twips):
            spacing.set(qn('w:after'), str(space_after_twips))
            changed += 1

        if text:
            next_text_para = para

    return changed


AW_TABLE_STYLE_IDS = frozenset({
    'AWStandardTable', 'AWComparisonTable', 'AWPhraseBankTable', 'AWRubricTable',
})


def apply_table_styles(doc, default_style: str = 'AW Standard Table',
                       space_after_twips: int = 280,
                       reference_doc_path=None) -> int:
    """Apply AW Standard Table style to all tables that have no custom style set.

    Also applies AW Table Header / AW Table Body paragraph styles to cell paragraphs
    so that font/size/spacing is set at the paragraph level (Word's table-style rPr
    is overridden by paragraph styles in the cascade).

    Copies AW Table Header and AW Table Body from the reference DOCX into the output
    if they are missing — Pandoc only propagates styles that appear in the source
    markdown, so these paragraph styles are never present unless explicitly copied.

    Also adds space-before on the paragraph immediately following each table.
    """
    # Ensure the cell paragraph styles exist in the output DOCX
    _ensure_styles_from_reference(
        doc, reference_doc_path, ['AW Table Header', 'AW Table Body']
    )
    table_style = _get_style_by_name_or_id(doc.styles, default_style)
    cell_body_style = _get_style_by_name_or_id(doc.styles, 'AW Table Body')
    changed = 0

    body = doc._body._body  # lxml element
    children = list(body)

    for idx, child in enumerate(children):
        # Apply default style to unstyled tables
        if child.tag == qn('w:tbl') and table_style:
            # Skip response placeholder tables (tagged with tblDescription)
            desc = child.find(f'.//{qn("w:tblDescription")}')
            if desc is not None and desc.get(qn('w:val')) == 'ResponsePlaceholder':
                pass
            else:
                from docx.table import Table as DTable
                try:
                    tbl = DTable(child, doc._body)
                    current = tbl.style.name if tbl.style else None
                    if current in (None, 'Table Normal', 'Normal Table'):
                        tbl.style = table_style
                        changed += 1

                    # Apply AW Table Body to all body-row cell paragraphs so the
                    # font/size/spacing is set at the paragraph level (Word's table-
                    # style rPr is overridden by paragraph styles in the cascade).
                    tbl_style_id = tbl.style.style_id if tbl.style else None
                    if tbl_style_id in AW_TABLE_STYLE_IDS:
                        cell_header_style = _get_style_by_name_or_id(doc.styles, 'AW Table Header')
                        rows = tbl.rows
                        for row_idx, row in enumerate(rows):
                            is_header = (row_idx == 0)
                            target = cell_header_style if is_header else cell_body_style
                            if target is None:
                                continue
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if para.style and para.style.name not in (
                                        'AW Table Body', 'AW Table Header'
                                    ):
                                        para.style = target

                    # Enforce 100% width and autofit directly on the table element
                    # (Pandoc sets explicit tblW which overrides style-level defaults)
                    tbl_pr = child.find(qn('w:tblPr'))
                    if tbl_pr is None:
                        tbl_pr = OxmlElement('w:tblPr')
                        child.insert(0, tbl_pr)
                    tbl_w = tbl_pr.find(qn('w:tblW'))
                    if tbl_w is None:
                        tbl_w = OxmlElement('w:tblW')
                        tbl_pr.append(tbl_w)
                    tbl_w.set(qn('w:w'), '5000')
                    tbl_w.set(qn('w:type'), 'pct')
                    tbl_lay = tbl_pr.find(qn('w:tblLayout'))
                    if tbl_lay is None:
                        tbl_lay = OxmlElement('w:tblLayout')
                        tbl_pr.append(tbl_lay)
                    tbl_lay.set(qn('w:type'), 'autofit')
                except Exception:
                    pass

            # Add space-before to the paragraph immediately following the table
            for next_child in children[idx + 1:]:
                if next_child.tag == qn('w:p'):
                    p_pr = next_child.find(qn('w:pPr'))
                    if p_pr is None:
                        p_pr = OxmlElement('w:pPr')
                        next_child.insert(0, p_pr)
                    spacing = p_pr.find(qn('w:spacing'))
                    if spacing is None:
                        spacing = OxmlElement('w:spacing')
                        p_pr.append(spacing)
                    if spacing.get(qn('w:before')) != str(space_after_twips):
                        spacing.set(qn('w:before'), str(space_after_twips))
                    break
                elif next_child.tag == qn('w:tbl'):
                    break  # next element is another table, skip

    return changed


def apply_response_placeholders(doc) -> int:
    """Replace {{PH-N: label}} marker paragraphs with ruled-line Word tables.

    Each placeholder is a full-width (fit-to-window) table with N ruled rows
    (bottom border only, fixed height, single line spacing). No header row —
    the writing instruction appears as normal text above the marker in the
    source document, so rows= directly equals the number of writing lines.

    Placeholder tables get a thin spacer paragraph after them so following
    prompts or list items do not run tight against the writing lines.

    Row counts per type (fallback defaults when rows= is absent):
      PH-1a:  1 row  (~0.8 cm)
      PH-1b:  2 rows (~1.6 cm)
      PH-1c:  3 rows (~2.4 cm)
      PH-1:   4 rows (~3.2 cm)
      PH-2:   6 rows (~4.8 cm)
      PH-3:  11 rows (~8.8 cm)
      PH-4:  19 rows (~15.2 cm)
      PH-5:  28 rows (~22.4 cm)
    """
    import re
    from lxml import etree

    PH_ROWS = {
        'PH-1a': 1, 'PH-1b': 2, 'PH-1c': 3,
        'PH-1': 4, 'PH-2': 6, 'PH-3': 11, 'PH-4': 19, 'PH-5': 28,
    }
    ROW_HEIGHT_TWIPS = 454      # ~0.8 cm per row
    RULE_COLOR = 'AAAAAA'
    SPACER_HEIGHT = '80'        # ~4pt spacer paragraph around response tables
    LIST_TEXT_INDENT_TWIPS = 357

    PH_RE = re.compile(r'^(.*?)\s*\{\{(PH-\d+[a-z]?):\s*([^}]+)\}\}\s*$')
    ROWS_RE = re.compile(r'\brows\s*=\s*(\d+)\b')
    # Tag to identify our placeholder tables so apply_table_styles skips them
    PH_TAG = 'ResponsePlaceholder'

    W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'

    def w(tag): return f'{{{W_NS}}}{tag}'
    def sa(el, attr, val): el.set(f'{{{W_NS}}}{attr}', val)

    body = doc._body._body

    def make_spacer_para():
        p = etree.Element(w('p'))
        pPr = etree.SubElement(p, w('pPr'))
        sp = etree.SubElement(pPr, w('spacing'))
        sa(sp, 'before', SPACER_HEIGHT)
        sa(sp, 'after', '0')
        sa(sp, 'line', '240')
        sa(sp, 'lineRule', 'auto')
        return p

    def parse_placeholder_payload(payload: str):
        parts = [part.strip() for part in payload.split('|')]
        id_label = parts[0]
        explicit_rows = None
        for part in parts[1:]:
            rows_match = ROWS_RE.search(part)
            if rows_match:
                explicit_rows = max(1, int(rows_match.group(1)))
                break
        return id_label, explicit_rows

    def is_list_paragraph(el) -> bool:
        if el is None or el.tag != qn('w:p'):
            return False
        pPr = el.find(w('pPr'))
        if pPr is None:
            return False
        if pPr.find(w('numPr')) is not None:
            return True
        p_style = pPr.find(w('pStyle'))
        if p_style is None:
            return False
        style_val = p_style.get(f'{{{W_NS}}}val') or ''
        return style_val.startswith(('ListNumber', 'ListBullet', 'Checklist'))

    def list_key(el):
        if not is_list_paragraph(el):
            return None
        pPr = el.find(w('pPr'))
        p_style = pPr.find(w('pStyle')) if pPr is not None else None
        numpr = pPr.find(w('numPr')) if pPr is not None else None
        numid = numpr.find(w('numId')) if numpr is not None else None
        ilvl = numpr.find(w('ilvl')) if numpr is not None else None
        return (
            p_style.get(f'{{{W_NS}}}val') if p_style is not None else '',
            numid.get(f'{{{W_NS}}}val') if numid is not None else '',
            ilvl.get(f'{{{W_NS}}}val') if ilvl is not None else '',
        )

    def is_checklist_paragraph(el) -> bool:
        if not is_list_paragraph(el):
            return False
        pPr = el.find(w('pPr'))
        p_style = pPr.find(w('pStyle')) if pPr is not None else None
        style_val = p_style.get(f'{{{W_NS}}}val') if p_style is not None else ''
        return style_val.startswith('Checklist')

    def is_ordered_list_paragraph(el) -> bool:
        if not is_list_paragraph(el):
            return False
        pPr = el.find(w('pPr'))
        p_style = pPr.find(w('pStyle')) if pPr is not None else None
        style_val = p_style.get(f'{{{W_NS}}}val') if p_style is not None else ''
        return style_val.startswith('ListNumber')

    def set_flush_number_list_indent(el) -> None:
        """Align list number to the margin and keep text on a hanging indent."""
        if not is_list_paragraph(el):
            return
        pPr = el.find(w('pPr'))
        if pPr is None:
            pPr = etree.SubElement(el, w('pPr'))
        ind = pPr.find(w('ind'))
        if ind is None:
            ind = etree.SubElement(pPr, w('ind'))
        sa(ind, 'left', str(LIST_TEXT_INDENT_TWIPS))
        sa(ind, 'hanging', str(LIST_TEXT_INDENT_TWIPS))

    def normalize_single_item_placeholder_list(body_children, marker_idx) -> bool:
        """Apply list/table alignment only when a placeholder follows one list item.

        Case 1: one placeholder after each list item. The table should align with
        that item's hanging text indent, and the item number should sit at the
        left margin.

        Case 2: one placeholder after a complete multi-item list. Leave the list
        and table in their normal positions.
        """
        if marker_idx <= 0:
            return False
        prev = body_children[marker_idx - 1]
        if not is_list_paragraph(prev):
            return False
        if not is_ordered_list_paragraph(prev) or is_checklist_paragraph(prev):
            return False
        key = list_key(prev)

        run_start = marker_idx - 1
        while (
            run_start - 1 >= 0
            and is_list_paragraph(body_children[run_start - 1])
            and list_key(body_children[run_start - 1]) == key
        ):
            run_start -= 1

        run_end = marker_idx - 1
        run_length = run_end - run_start + 1
        if run_length != 1:
            return False

        set_flush_number_list_indent(prev)
        return True

    def text_width_twips() -> int:
        sectPr = body.find(w('sectPr'))
        if sectPr is None:
            return 9360
        pgSz = sectPr.find(w('pgSz'))
        pgMar = sectPr.find(w('pgMar'))
        if pgSz is None or pgMar is None:
            return 9360
        page_w = int(pgSz.get(f'{{{W_NS}}}w') or 12240)
        left = int(pgMar.get(f'{{{W_NS}}}left') or 1440)
        right = int(pgMar.get(f'{{{W_NS}}}right') or 1440)
        return max(1000, page_w - left - right)

    TEXT_WIDTH_TWIPS = text_width_twips()

    def make_tbl(n_rows, left_indent_twips: int = 0):
        tbl = etree.Element(w('tbl'))

        # tblPr — fit to window (pct type, 5000 = 100%)
        tbl_pr = etree.SubElement(tbl, w('tblPr'))
        # Mark as placeholder so apply_table_styles skips it
        tbl_desc = etree.SubElement(tbl_pr, w('tblDescription'))
        sa(tbl_desc, 'val', PH_TAG)
        if left_indent_twips:
            tbl_ind = etree.SubElement(tbl_pr, w('tblInd'))
            sa(tbl_ind, 'w', str(left_indent_twips)); sa(tbl_ind, 'type', 'dxa')
        tbl_w = etree.SubElement(tbl_pr, w('tblW'))
        if left_indent_twips:
            sa(tbl_w, 'w', str(max(1000, TEXT_WIDTH_TWIPS - left_indent_twips)))
            sa(tbl_w, 'type', 'dxa')
        else:
            sa(tbl_w, 'w', '5000'); sa(tbl_w, 'type', 'pct')  # 100% of text width
        tbl_layout = etree.SubElement(tbl_pr, w('tblLayout'))
        sa(tbl_layout, 'type', 'fixed')
        # No outer borders
        tbl_borders = etree.SubElement(tbl_pr, w('tblBorders'))
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            el = etree.SubElement(tbl_borders, w(side))
            sa(el, 'val', 'none'); sa(el, 'sz', '0'); sa(el, 'space', '0'); sa(el, 'color', 'auto')
        # No spacing before/after the table itself
        tbl_look = etree.SubElement(tbl_pr, w('tblLook'))
        sa(tbl_look, 'val', '0000')

        def make_row():
            tr = etree.Element(w('tr'))
            trPr = etree.SubElement(tr, w('trPr'))
            trH = etree.SubElement(trPr, w('trHeight'))
            sa(trH, 'val', str(ROW_HEIGHT_TWIPS)); sa(trH, 'hRule', 'exact')

            tc = etree.SubElement(tr, w('tc'))
            tcPr = etree.SubElement(tc, w('tcPr'))
            tcW = etree.SubElement(tcPr, w('tcW'))
            sa(tcW, 'w', '5000'); sa(tcW, 'type', 'pct')
            # Vertical centering
            vAlign = etree.SubElement(tcPr, w('vAlign'))
            sa(vAlign, 'val', 'center')
            # Cell borders: only bottom ruled line
            tc_borders = etree.SubElement(tcPr, w('tcBorders'))
            for side in ('top', 'left', 'right'):
                el = etree.SubElement(tc_borders, w(side))
                sa(el, 'val', 'none'); sa(el, 'sz', '0'); sa(el, 'space', '0'); sa(el, 'color', 'auto')
            bot = etree.SubElement(tc_borders, w('bottom'))
            sa(bot, 'val', 'single'); sa(bot, 'sz', '4')
            sa(bot, 'space', '0'); sa(bot, 'color', RULE_COLOR)

            p = etree.SubElement(tc, w('p'))
            pPr = etree.SubElement(p, w('pPr'))
            sp = etree.SubElement(pPr, w('spacing'))
            sa(sp, 'before', '0'); sa(sp, 'after', '0')
            sa(sp, 'line', '240'); sa(sp, 'lineRule', 'auto')
            jc = etree.SubElement(pPr, w('jc')); sa(jc, 'val', 'left')
            return tr

        for _ in range(n_rows):
            tbl.append(make_row())
        return tbl

    # First pass: collect marker paragraphs
    markers = []
    for child in list(body):
        if child.tag != qn('w:p'):
            continue
        text = ''.join(t.text or '' for t in child.iter(qn('w:t'))).strip()
        m = PH_RE.match(text)
        if not m:
            continue
        prefix_label = m.group(1).strip()
        ph_type = m.group(2)
        id_label, explicit_rows = parse_placeholder_payload(m.group(3).strip())
        if ph_type not in PH_ROWS:
            continue
        if prefix_label:
            display_label = prefix_label
        else:
            parts = id_label.split('-')
            display_label = ' '.join(p.capitalize() for p in parts[2:]) if len(parts) > 2 else id_label
        markers.append((child, ph_type, display_label, explicit_rows))

    # Second pass: replace in reverse order so indices stay valid
    changed = 0
    for child, ph_type, display_label, explicit_rows in reversed(markers):
        n_rows = explicit_rows or PH_ROWS[ph_type]
        body_children = list(body)
        idx = body_children.index(child)
        follows_list = normalize_single_item_placeholder_list(body_children, idx)
        tbl_el = make_tbl(n_rows, LIST_TEXT_INDENT_TWIPS if follows_list else 0)
        body.remove(child)
        body.insert(idx, tbl_el)
        changed += 1

    # Third pass: insert spacer paragraphs after placeholder tables
    children = list(body)
    inserts = []
    for i, child in enumerate(children):
        if child.tag != qn('w:tbl'):
            continue
        desc = child.find(f'.//{w("tblDescription")}')
        if desc is None or desc.get(f'{{{W_NS}}}val') != PH_TAG:
            continue
        if i + 1 >= len(children):
            continue
        nxt = children[i + 1]
        if nxt.tag == qn('w:sectPr'):
            continue
        inserts.append((i + 1, make_spacer_para()))

    # Insert spacers in reverse so positions stay correct
    for pos, spacer in reversed(inserts):
        body.insert(pos, spacer)

    return changed


def insert_section_after_toc(
    docx_path,
    has_toc=True,
    insert_h1_sections=True,
    reference_doc_path=None,
    apply_semantic_labels=False,
    building_block_template=None,
    tag_style: str = DIV_TAG_ICON_STYLE,
    profile: bool = False,
    progress: bool = True,
    progress_warn_seconds: float = 30.0,
):
    """
    Post-process `docx_path` to:
    1. If has_toc: insert section break after TOC, restart page numbering at 1,
       remove page numbers from TOC
    2. Insert section breaks before H1 headings
    3. Apply semantic label rendering when apply_semantic_labels=True
    4. Insert Quick Parts with Word COM when available
    """
    profiler = PostprocessProfiler(
        profile,
        progress=progress,
        warn_after_seconds=progress_warn_seconds,
    )
    context = PostprocessContext(reference_doc_path)
    docx_path = Path(docx_path)

    def save_and_reopen(doc_to_save, name: str):
        with profiler.step(name) as timing:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=docx_path.parent)
            os.close(tmp_fd)
            try:
                doc_to_save.save(tmp_path)
                shutil.move(tmp_path, docx_path)
                timing['status'] = 'saved'
                return Document(docx_path)
            except Exception:
                try:
                    os.unlink(tmp_path)
                except OSError:
                    pass
                raise

    def run_pass(name: str, func, message: str | None = None, *, enabled: bool = True) -> int:
        nonlocal made_change
        if not enabled:
            profiler.skip(name)
            return 0
        with profiler.step(name) as timing:
            count = func()
            timing['changed'] = count
            timing['status'] = 'changed' if count else 'no-change'
        if count > 0:
            if message:
                print(message.format(count=count))
            made_change = True
        return count

    with profiler.step('load docx'):
        doc = Document(docx_path)
    paragraphs = list(doc.paragraphs)

    toc_par_index = None
    if has_toc:
        with profiler.step('scan toc') as timing:
            for index, paragraph in enumerate(paragraphs):
                fld_simples = paragraph._p.findall('.//w:fldSimple', NS)
                for fld in fld_simples:
                    instr = fld.get(qn('w:instr')) or fld.get('instr')
                    if instr and instr.strip().upper().startswith('TOC'):
                        toc_par_index = index
                fld_chars = paragraph._p.findall('.//w:fldChar', NS)
                if fld_chars:
                    instr_texts = paragraph._p.findall('.//w:instrText', NS)
                    for instr_text in instr_texts:
                        if instr_text.text and instr_text.text.strip().upper().startswith('TOC'):
                            toc_par_index = index
            timing['changed'] = 1 if toc_par_index is not None else 0
            timing['status'] = 'found' if toc_par_index is not None else 'not-found'
    else:
        profiler.skip('scan toc')

    made_change = False

    if toc_par_index is not None:
        toc_par = paragraphs[toc_par_index]
        with profiler.step('toc section break') as timing:
            try:
                _apply_next_page_section_to_paragraph(toc_par, start_page=1)
                timing['changed'] = 1
                timing['status'] = 'changed'
                made_change = True
                print('Inserted section break after TOC with page numbering restart at 1')
            except Exception as exc:
                timing['changed'] = 0
                timing['status'] = 'warning'
                print(f'Warning: Could not insert section break after TOC: {exc}')

        try:
            doc = save_and_reopen(doc, 'save/reopen after toc')
            if len(doc.sections) > 0:
                toc_section = doc.sections[0]
                with profiler.step('remove toc page numbers') as timing:
                    _remove_page_numbers_from_footer(toc_section.footer)
                    try:
                        _remove_page_numbers_from_footer(toc_section.first_page_footer)
                    except Exception:
                        pass
                    timing['changed'] = 1
                    timing['status'] = 'changed'
                print('Removed page numbers from TOC section')
                made_change = True
        except Exception as exc:
            print(f'Warning: Could not remove page numbers from TOC: {exc}')
    else:
        profiler.skip('toc section break')
        profiler.skip('save/reopen after toc')
        profiler.skip('remove toc page numbers')

    run_pass(
        'normalize module headings',
        lambda: normalize_module_headings(doc),
        'Normalized {count} module heading item(s)',
    )

    def h1_sections_pass() -> int:
        try:
            return insert_section_breaks_before_h1(
                doc,
                skip_first=True,
                restart_first=False,
            )
        except Exception as exc:
            print(f'Warning: Could not insert H1 section breaks: {exc}')
            return 0

    run_pass(
        'h1 section breaks',
        h1_sections_pass,
        'Inserted {count} section break(s) before H1 headings',
        enabled=insert_h1_sections,
    )

    run_pass('list styles', lambda: apply_list_styles(doc), 'Applied list styles to {count} paragraph(s)')
    run_pass(
        'strip alpha markers',
        lambda: strip_literal_alpha_markers(doc),
        'Removed literal alphabetic markers from {count} list paragraph(s)',
    )
    run_pass('checklist style', lambda: apply_checklist_style(doc), 'Applied Checklist style to {count} item(s)')
    profiler.skip('example block styles (source-driven)')
    run_pass(
        'post-list spacing',
        lambda: apply_spacing_after_lists(doc),
        'Applied post-list spacing to {count} paragraph(s)',
    )
    run_pass(
        'table styles',
        lambda: apply_table_styles(doc, reference_doc_path=reference_doc_path),
        'Applied table styles to {count} table(s)',
    )
    run_pass(
        'response placeholders',
        lambda: apply_response_placeholders(doc),
        'Inserted {count} response placeholder(s)',
    )
    profiler.skip('semantic paragraph styles (source-driven)')
    run_pass(
        'semantic div labels',
        lambda: apply_semantic_div_labels(doc, reference_doc_path, tag_style, context),
        'Updated semantic Div labels in {count} place(s)',
        enabled=apply_semantic_labels,
    )
    run_pass(
        'body text normalization',
        lambda: apply_body_text_to_normal_paragraphs(doc),
        'Applied Body Text to {count} normal paragraph(s)',
    )
    run_pass(
        'post-example spacing',
        lambda: apply_spacing_after_example_blocks(doc),
        'Applied post-example spacing to {count} paragraph(s)',
    )
    run_pass(
        'strip heading activity codes',
        lambda: strip_activity_codes_from_headings(doc),
        'Removed activity codes from {count} heading(s)',
    )
    run_pass(
        'remove pandoc fallback styles',
        lambda: remove_pandoc_generated_styles(doc),
        'Replaced Pandoc fallback styles in {count} place(s)',
    )
    run_pass(
        'purge non-reference styles',
        lambda: purge_styles_not_in_reference(doc, reference_doc_path, context),
        'Removed non-reference style usage/definitions in {count} place(s)',
    )
    run_pass(
        'disable heading page breaks',
        lambda: disable_heading_style_page_breaks(doc),
        'Disabled page-break-before on {count} heading style(s)',
    )
    run_pass(
        'module/unit section breaks',
        lambda: insert_section_breaks_before_modules_and_units(doc),
        'Inserted {count} module/unit section break(s)',
    )

    if made_change:
        doc = save_and_reopen(doc, 'save/reopen after main passes')
    else:
        profiler.skip('save/reopen after main passes')

    final_change = False
    header_changes = run_pass(
        'running headers',
        lambda: update_running_headers(doc),
        'Updated running headers in {count} place(s)',
    )
    if header_changes > 0:
        final_change = True

    footer_changes = run_pass(
        'section footer references',
        lambda: ensure_section_footer_references(doc, skip_first_section=has_toc),
        'Restored section footer references in {count} place(s)',
    )
    if footer_changes > 0:
        final_change = True

    # Unit title table substitution — runs whenever reference_doc is available,
    # independent of apply_semantic_labels.
    unit_tile_changes = run_pass(
        'unit title tables',
        lambda: replace_unit_headings_with_title_tables(
            doc,
            reference_doc_path=reference_doc_path,
            context=context,
        ),
        'Inserted {count} unit title table(s)',
    )
    if unit_tile_changes > 0:
        final_change = True

    restored_headings = run_pass(
        'restore unit overview headings',
        lambda: restore_unit_overview_headings_after_unit_tiles(doc),
        'Restored {count} Unit Overview heading(s) after unit tiles',
        enabled=unit_tile_changes > 0,
    )
    if restored_headings > 0:
        final_change = True

    if final_change:
        doc = save_and_reopen(doc, 'save/reopen final changes')
        made_change = True
    else:
        profiler.skip('save/reopen final changes')

    profiler.summary()
    return made_change


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument('docx', help='DOCX file to postprocess')
    parser.add_argument('--toc', action='store_true', default=False, help='Document has a Table of Contents (opt-in)')
    parser.add_argument(
        '--reference-doc',
        help='Reference DOCX used as the canonical source for styles and document properties.',
    )
    parser.add_argument(
        '--building-block-template',
        help='Word template containing BuildingBlockEntries used for Quick Part insertion.',
    )
    parser.add_argument(
        '--h1-sections',
        action='store_true',
        default=False,
        help='Insert nextPage section breaks before H1 headings (opt-in; use when the styleref does not set pageBreakBefore on Heading 1)',
    )
    parser.add_argument(
        '--apply-semantic-labels',
        action='store_true',
        help=(
            'Apply semantic Div label rendering: insert icons, set label character styles, '
            'normalize label text, and replace unit headings with title tables (Phase C). '
            'Off by default — structural cleanup runs regardless.'
        ),
    )
    parser.add_argument(
        '--no-semantic-formatting',
        action='store_true',
        help='Deprecated. Semantic formatting is off by default; use --apply-semantic-labels to enable it.',
    )
    parser.add_argument(
        '--tag-style',
        choices=['filled', 'outline'],
        default=DIV_TAG_ICON_STYLE,
        help='Which tag icon variant to insert before Div label text (default: filled).',
    )
    parser.add_argument(
        '--profile',
        action='store_true',
        help='Print timing information for each DOCX postprocess pass and save/reopen step.',
    )
    parser.add_argument(
        '--progress-warn-seconds',
        type=float,
        default=30.0,
        help='Print a warning when a postprocess pass runs longer than this many seconds (default: 30; use 0 to disable warnings).',
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)

    did_update = insert_section_after_toc(
        args.docx,
        has_toc=args.toc,
        insert_h1_sections=args.h1_sections,
        reference_doc_path=args.reference_doc,
        apply_semantic_labels=args.apply_semantic_labels,
        building_block_template=args.building_block_template,
        tag_style=args.tag_style,
        profile=args.profile,
        progress=True,
        progress_warn_seconds=args.progress_warn_seconds,
    )

    if did_update:
        print('Post-processing complete.')
    else:
        print('No changes made.')
    return 0


if __name__ == '__main__':
    sys.exit(main())

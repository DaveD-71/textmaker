"""
Preprocess a DOCX before pandoc conversion to preserve elements pandoc drops.

Approach:
- Mark page/section breaks with sentinel paragraphs.
- Preserve manual line breaks inside paragraphs.
- Expand common fields (REF, PAGEREF, HYPERLINK) into textual markers.
- Add placeholders for shapes/textboxes.

Sentinels are simple text markers that survive pandoc and are later replaced:
- [[PAGEBREAK]]
- [[SECTIONBREAK]]
- [[LINEBREAK]] inside paragraphs
- [[REF:id|label]] for cross-references
- [[SHAPE:index|alt text]] for shapes/text boxes (alt text optional)
"""
from __future__ import annotations

from pathlib import Path
from typing import List

try:
    from docx import Document  # type: ignore[reportMissingImports]
    from docx.oxml import OxmlElement  # type: ignore[reportMissingImports]
    from docx.oxml.ns import qn  # type: ignore[reportMissingImports]
    from docx.text.paragraph import Paragraph  # type: ignore[reportMissingImports]
except ImportError as exc:
    raise RuntimeError('python-docx is required. Install with `pip install python-docx`.') from exc


NS = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'v': 'urn:schemas-microsoft-com:vml',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
}


def _add_paragraph_after(paragraph, text: str):
    """Insert a paragraph immediately after the given paragraph."""
    new_p = OxmlElement('w:p')
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def _paragraph_has_section_break(paragraph) -> bool:
    p_pr = paragraph._p.find(qn('w:pPr'))
    if p_pr is None:
        return False
    sect = p_pr.find(qn('w:sectPr'))
    if sect is None:
        return False
    type_el = sect.find(qn('w:type'))
    if type_el is None:
        return True  # default is nextPage
    val = type_el.get(qn('w:val'))
    return val in (None, 'nextPage', 'continuous', 'oddPage', 'evenPage')


def _paragraph_has_page_break(paragraph) -> bool:
    # Look for w:br type="page" in runs
    for run in paragraph.runs:
        brs = run._r.findall(qn('w:br'))
        for br in brs:
            if br.get(qn('w:type')) == 'page':
                return True
    return False


def _replace_line_breaks(paragraph):
    """Replace line breaks in runs with [[LINEBREAK]] text tokens."""
    for run in paragraph.runs:
        new_elems: List = []
        for child in list(run._r):
            if child.tag == qn('w:br') and child.get(qn('w:type')) != 'page':
                new_text = OxmlElement('w:t')
                new_text.text = '[[LINEBREAK]]'
                new_elems.append(new_text)
            else:
                new_elems.append(child)
        # Rebuild run children
        for child in list(run._r):
            run._r.remove(child)
        for child in new_elems:
            run._r.append(child)


def _expand_fields(paragraph):
    """Expand common field codes to textual markers."""
    fld_simp = paragraph._p.findall('.//w:fldSimple', NS)
    for fld in fld_simp:
        instr = fld.get(qn('w:instr')) or fld.get('instr') or ''
        instr_up = instr.upper()
        if instr_up.startswith('HYPERLINK '):
            # Leave hyperlinks for pandoc; skip
            continue
        label = instr.replace('"', '').strip()
        marker = f'[[REF:{label}]]'
        paragraph.add_run(marker)

    # Complex fields: gather instrText
    instr_texts = paragraph._p.findall('.//w:instrText', NS)
    for instr_el in instr_texts:
        instr = (instr_el.text or '').strip()
        if not instr:
            continue
        instr_up = instr.upper()
        if instr_up.startswith('REF ') or instr_up.startswith('PAGEREF '):
            ref_id = instr.split()[1] if len(instr.split()) > 1 else instr
            marker = f'[[REF:{ref_id}]]'
            paragraph.add_run(marker)


def _shape_alt_text(shape_elem) -> str:
    """Extract human-readable shape label from DrawingML or legacy VML."""
    doc_prs = shape_elem.findall('.//wp:docPr', NS)
    if doc_prs:
        desc = (doc_prs[0].get('descr') or '').strip()
        title = (doc_prs[0].get('title') or '').strip()
        if desc or title:
            return desc or title

    v_shapes = shape_elem.findall('.//v:shape', NS)
    if v_shapes:
        attrs = v_shapes[0].attrib
        for key in ('alt', 'title', 'alttext'):
            val = (attrs.get(key) or '').strip()
            if val:
                return val

    return ''


def _add_shape_placeholders(paragraph, start_index: int) -> int:
    """
    Add placeholders for drawing elements that might otherwise be dropped.

    Marker IDs are deterministic and aligned with docx_to_markdown.extract_shapes().
    """
    shape_elements = paragraph._p.findall('.//w:drawing', NS) + paragraph._p.findall('.//w:pict', NS)
    shape_index = start_index

    for shape_elem in shape_elements:
        shape_index += 1
        alt = _shape_alt_text(shape_elem)
        if alt:
            paragraph.add_run(f'[[SHAPE:{shape_index}|{alt}]]')
        else:
            paragraph.add_run(f'[[SHAPE:{shape_index}]]')
    return shape_index


def preprocess_docx(src: Path, dest: Path) -> Path:
    """Write a preprocessed copy of src to dest with sentinel markers inserted."""
    doc = Document(src)
    shape_index = 0

    for p in list(doc._element.findall('.//w:p', NS)):
        paragraph = Paragraph(p, doc)
        # Mark section breaks and page breaks
        has_section = _paragraph_has_section_break(paragraph)
        has_page_break = _paragraph_has_page_break(paragraph)

        _replace_line_breaks(paragraph)
        _expand_fields(paragraph)
        shape_index = _add_shape_placeholders(paragraph, shape_index)

        if has_section:
            _add_paragraph_after(paragraph, '[[SECTIONBREAK]]')
        elif has_page_break:
            _add_paragraph_after(paragraph, '[[PAGEBREAK]]')

    doc.save(dest)
    return dest


def main() -> None:
    import argparse

    parser = argparse.ArgumentParser(
        description='Insert sentinel markers into a DOCX before pandoc.'
    )
    parser.add_argument('input', help='Source DOCX')
    parser.add_argument('output', help='Destination DOCX with markers')
    args = parser.parse_args()

    preprocess_docx(Path(args.input), Path(args.output))
    print(f'Wrote preprocessed DOCX to {args.output}')


if __name__ == '__main__':
    main()

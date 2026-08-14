"""
Generate a `reference.docx` with basic modern professional styles.
This uses `python-docx` to set Normal and Heading styles and A4 page size.
Run default: python scripts/generate_reference_docx.py --out reference.docx
Run from source: python scripts/generate_reference_docx.py --input source.docx --out reference.docx
Run from YAML: python scripts/generate_reference_docx.py --spec styles.yaml --out reference.docx
"""
# pylint: disable=protected-access,broad-exception-caught
import argparse
import sys
from pathlib import Path
try:
    # type: ignore[reportMissingImports]
    from docx import Document
    # type: ignore[reportMissingImports]
    from docx.shared import Mm, Pt, RGBColor
    # type: ignore[reportMissingImports]
    from docx.enum.style import WD_STYLE_TYPE
    # type: ignore[reportMissingImports]
    from docx.oxml import OxmlElement
    # type: ignore[reportMissingImports]
    from docx.oxml.ns import qn
    # type: ignore[reportMissingImports]
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    # type: ignore[reportMissingImports]
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ModuleNotFoundError as _err:
    # Only convert missing-docx into a helpful message; bubble up other missing deps.
    missing_name = getattr(_err, 'name', '') or ''
    if isinstance(missing_name, str) and missing_name.startswith('docx'):
        raise RuntimeError(
            "Missing dependency: python-docx is required. Install with `pip install python-docx` and retry."
        ) from _err
    raise
except ImportError as _err:
    # For other import issues, expose the original message so environment problems surface.
    raise RuntimeError(f"Failed to import python-docx dependency: {_err}") from _err
    # end of import guard


def _load_yaml(path: Path) -> dict:
    try:
        import yaml  # type: ignore[reportMissingImports]
    except ModuleNotFoundError as exc:
        raise RuntimeError(
            "Missing dependency: PyYAML is required for --spec. Install with `pip install pyyaml`."
        ) from exc
    with path.open(encoding='utf-8') as f:
        data = yaml.safe_load(f)
    if not isinstance(data, dict):
        raise ValueError(f'Style spec must be a YAML mapping: {path}')
    return data


def _clean_hex(value) -> str | None:
    if value is None:
        return None
    text = str(value).strip()
    if not text:
        return None
    if text.startswith('#'):
        text = text[1:]
    if len(text) != 6:
        raise ValueError(f'Expected 6-digit hex color, got: {value!r}')
    int(text, 16)
    return text.upper()


def _resolve_alias(value, aliases: dict, *, kind: str):
    if value is None:
        return None
    if isinstance(value, str) and value in aliases:
        return aliases[value]
    if kind == 'color':
        return _clean_hex(value)
    return value


def _rgb_color(value) -> RGBColor:
    hex_value = _clean_hex(value)
    if hex_value is None:
        raise ValueError('Color value is required')
    return RGBColor(
        int(hex_value[0:2], 16),
        int(hex_value[2:4], 16),
        int(hex_value[4:6], 16),
    )


def _length_pt(value):
    return None if value is None else Pt(float(value))


def _length_mm(value):
    return None if value is None else Mm(float(value))


def _set_style_rfonts(style, font_name: str):
    """Ensure explicit run fonts override theme fonts."""
    try:
        el = style._element
        rPr = el.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            el.insert(0, rPr)
        for existing in list(rPr.findall(qn('w:rFonts'))):
            rPr.remove(existing)
        rfonts = OxmlElement('w:rFonts')
        rfonts.set(qn('w:ascii'), font_name)
        rfonts.set(qn('w:hAnsi'), font_name)
        rfonts.set(qn('w:cs'), font_name)
        rfonts.set(qn('w:eastAsia'), font_name)
        rPr.append(rfonts)
    except Exception:
        # best-effort; ignore if style internals differ
        pass


def _ensure_run_properties(style):
    el = style._element
    r_pr = el.find(qn('w:rPr'))
    if r_pr is None:
        r_pr = OxmlElement('w:rPr')
        el.insert(0, r_pr)
    return r_pr


def _ensure_paragraph_properties(style):
    el = style._element
    p_pr = el.find(qn('w:pPr'))
    if p_pr is None:
        p_pr = OxmlElement('w:pPr')
        el.insert(0, p_pr)
    return p_pr


def _set_style_language(style, lang: str):
    r_pr = _ensure_run_properties(style)
    for existing in list(r_pr.findall(qn('w:lang'))):
        r_pr.remove(existing)
    lang_el = OxmlElement('w:lang')
    lang_el.set(qn('w:val'), lang)
    lang_el.set(qn('w:eastAsia'), lang)
    lang_el.set(qn('w:bidi'), lang)
    r_pr.append(lang_el)


def _set_on_off_ppr(style, tag: str, enabled: bool):
    p_pr = _ensure_paragraph_properties(style)
    for existing in list(p_pr.findall(qn(f'w:{tag}'))):
        p_pr.remove(existing)
    el = OxmlElement(f'w:{tag}')
    el.set(qn('w:val'), '1' if enabled else '0')
    p_pr.append(el)


def _set_on_off_settings(doc, tag: str, enabled: bool):
    settings = doc.settings._element
    for existing in list(settings.findall(qn(f'w:{tag}'))):
        settings.remove(existing)
    el = OxmlElement(f'w:{tag}')
    el.set(qn('w:val'), '1' if enabled else '0')
    settings.append(el)


def _set_p_borders(style, borders):
    """Apply paragraph borders to a style using w:pBdr."""
    p_pr = _ensure_paragraph_properties(style)
    for existing in list(p_pr.findall(qn('w:pBdr'))):
        p_pr.remove(existing)
    if not borders:
        return
    p_bdr = OxmlElement('w:pBdr')
    for side, attrs in borders.items():
        side_el = OxmlElement(f'w:{side}')
        for key, val in attrs.items():
            attr_val = _clean_hex(val) if key == 'color' else val
            side_el.set(qn(f'w:{key}'), str(attr_val))
        p_bdr.append(side_el)
    p_pr.append(p_bdr)


def _set_shading(style, fill=None, color="auto", val="clear", theme_fill=None, theme_shade=None):
    """Apply paragraph shading to a style; pass fill=None to clear shading."""
    p_pr = _ensure_paragraph_properties(style)
    for existing in list(p_pr.findall(qn('w:shd'))):
        p_pr.remove(existing)
    if fill is None:
        return
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), val)
    shd.set(qn('w:color'), color)
    shd.set(qn('w:fill'), _clean_hex(fill))
    if theme_fill:
        shd.set(qn('w:themeFill'), theme_fill)
    if theme_shade:
        shd.set(qn('w:themeFillShade'), theme_shade)
    p_pr.append(shd)


def _add_contextual_spacing(style):
    """Add w:contextualSpacing to a style's paragraph properties."""
    p_pr = _ensure_paragraph_properties(style)
    if p_pr.find(qn('w:contextualSpacing')) is None:
        p_pr.append(OxmlElement('w:contextualSpacing'))


def _set_keep_with_next(style, enabled: bool):
    _set_on_off_ppr(style, 'keepNext', enabled)


def _set_snap_to_grid(style, enabled: bool):
    _set_on_off_ppr(style, 'snapToGrid', enabled)


def _set_table_cell_margins(style, margins: dict):
    tbl_pr = style._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        style._element.insert(0, tbl_pr)
    for existing in list(tbl_pr.findall(qn('w:tblCellMar'))):
        tbl_pr.remove(existing)
    cell_mar = OxmlElement('w:tblCellMar')
    for side in ('top', 'left', 'bottom', 'right'):
        if side not in margins:
            continue
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:w'), str(int(round(float(margins[side]) * 56.7))))  # mm -> twips
        el.set(qn('w:type'), 'dxa')
        cell_mar.append(el)
    tbl_pr.append(cell_mar)


def _set_table_borders(style, borders: dict):
    tbl_pr = style._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        style._element.insert(0, tbl_pr)
    for existing in list(tbl_pr.findall(qn('w:tblBorders'))):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement('w:tblBorders')
    for side, attrs in borders.items():
        el = OxmlElement(f'w:{side}')
        for key, val in attrs.items():
            attr_val = _clean_hex(val) if key == 'color' else val
            el.set(qn(f'w:{key}'), str(attr_val))
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


def _apply_font(style, spec: dict, *, fonts: dict, colors: dict):
    font_name = _resolve_alias(spec.get('font'), fonts, kind='font')
    if font_name:
        style.font.name = str(font_name)
        _set_style_rfonts(style, str(font_name))
    if spec.get('size_pt') is not None:
        style.font.size = Pt(float(spec['size_pt']))
    if spec.get('bold') is not None:
        style.font.bold = bool(spec['bold'])
    if spec.get('italic') is not None:
        style.font.italic = bool(spec['italic'])
    if spec.get('underline') is not None:
        style.font.underline = bool(spec['underline'])
    color = _resolve_alias(spec.get('color'), colors, kind='color')
    if color:
        style.font.color.rgb = _rgb_color(color)
    if spec.get('language') is not None:
        _set_style_language(style, str(spec['language']))


def _apply_paragraph_format(style, spec: dict):
    p = style.paragraph_format
    if spec.get('space_before_pt') is not None:
        p.space_before = Pt(float(spec['space_before_pt']))
    if spec.get('space_after_pt') is not None:
        p.space_after = Pt(float(spec['space_after_pt']))
    if spec.get('line_spacing') is not None:
        p.line_spacing = float(spec['line_spacing'])
    if spec.get('left_indent_mm') is not None:
        p.left_indent = Mm(float(spec['left_indent_mm']))
    if spec.get('right_indent_mm') is not None:
        p.right_indent = Mm(float(spec['right_indent_mm']))
    if spec.get('first_line_indent_mm') is not None:
        p.first_line_indent = Mm(float(spec['first_line_indent_mm']))
    if spec.get('page_break_before') is not None:
        p.page_break_before = bool(spec['page_break_before'])
    if spec.get('keep_with_next') is not None:
        _set_keep_with_next(style, bool(spec['keep_with_next']))
    if spec.get('widow_control') is not None:
        _set_on_off_ppr(style, 'widowControl', bool(spec['widow_control']))
    if spec.get('suppress_auto_hyphens') is not None:
        _set_on_off_ppr(style, 'suppressAutoHyphens', bool(spec['suppress_auto_hyphens']))
    if spec.get('snap_to_grid') is not None:
        _set_snap_to_grid(style, bool(spec['snap_to_grid']))
    if spec.get('contextual_spacing'):
        _add_contextual_spacing(style)
    alignment = spec.get('alignment')
    if alignment:
        align_map = {
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'centre': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
            'justify': WD_PARAGRAPH_ALIGNMENT.JUSTIFY,
        }
        p.alignment = align_map[str(alignment).lower()]


def _get_or_add_style(styles, name: str, style_type):
    if name in styles:
        return styles[name]
    return styles.add_style(name, style_type)


def _apply_style_spec(doc, style_spec: dict, *, fonts: dict, colors: dict):
    styles = doc.styles
    name = style_spec.get('name')
    if not name:
        raise ValueError(f'Style entry is missing name: {style_spec!r}')
    style_type_name = str(style_spec.get('type', 'paragraph')).lower()
    type_map = {
        'paragraph': WD_STYLE_TYPE.PARAGRAPH,
        'character': WD_STYLE_TYPE.CHARACTER,
        'table': WD_STYLE_TYPE.TABLE,
    }
    if style_type_name not in type_map:
        raise ValueError(f"Unsupported style type for {name!r}: {style_type_name!r}")
    style = _get_or_add_style(styles, str(name), type_map[style_type_name])

    base = style_spec.get('base')
    if base and base in styles and style_type_name != 'table':
        style.base_style = styles[base]

    if style_type_name in {'paragraph', 'character'}:
        _apply_font(style, style_spec, fonts=fonts, colors=colors)
    if style_type_name == 'paragraph':
        _apply_paragraph_format(style, style_spec.get('paragraph', style_spec))
        shading = _resolve_alias(style_spec.get('shading'), colors, kind='color')
        if shading:
            _set_shading(style, fill=shading)
        borders = style_spec.get('borders')
        if borders:
            resolved_borders = {}
            for side, attrs in borders.items():
                resolved_borders[side] = {
                    key: _resolve_alias(val, colors, kind='color') if key == 'color' else val
                    for key, val in attrs.items()
                }
            _set_p_borders(style, resolved_borders)
    if style_type_name == 'table':
        table_spec = style_spec.get('table', {})
        if table_spec.get('alignment'):
            align_map = {
                'left': WD_TABLE_ALIGNMENT.LEFT,
                'center': WD_TABLE_ALIGNMENT.CENTER,
                'centre': WD_TABLE_ALIGNMENT.CENTER,
                'right': WD_TABLE_ALIGNMENT.RIGHT,
            }
            style.paragraph_format.alignment = None
            # python-docx does not expose table-style alignment directly; keep XML props below.
        if table_spec.get('cell_margins_mm'):
            _set_table_cell_margins(style, table_spec['cell_margins_mm'])
        if table_spec.get('borders'):
            resolved_borders = {}
            for side, attrs in table_spec['borders'].items():
                resolved_borders[side] = {
                    key: _resolve_alias(val, colors, kind='color') if key == 'color' else val
                    for key, val in attrs.items()
                }
            _set_table_borders(style, resolved_borders)
    return style


def create_reference_from_spec(spec_path: str, out_path: str):
    """Create a reference DOCX from a YAML style specification."""
    spec = _load_yaml(Path(spec_path))
    doc = Document()
    section = doc.sections[0]

    page = spec.get('page', {})
    size = str(page.get('size', 'A4')).lower()
    if size == 'a4':
        section.page_height = Mm(297)
        section.page_width = Mm(210)
    else:
        section.page_width = _length_mm(page.get('width_mm')) or section.page_width
        section.page_height = _length_mm(page.get('height_mm')) or section.page_height

    margins = page.get('margins_mm', {})
    if isinstance(margins, (int, float)):
        margins = {'top': margins, 'bottom': margins, 'left': margins, 'right': margins}
    if margins:
        if margins.get('top') is not None:
            section.top_margin = Mm(float(margins['top']))
        if margins.get('bottom') is not None:
            section.bottom_margin = Mm(float(margins['bottom']))
        if margins.get('left') is not None:
            section.left_margin = Mm(float(margins['left']))
        if margins.get('right') is not None:
            section.right_margin = Mm(float(margins['right']))

    fonts = spec.get('fonts', {})
    colors = {name: _clean_hex(value) for name, value in spec.get('colors', {}).items()}

    document_settings = spec.get('document', {})
    if document_settings.get('hyphenation') is not None:
        _set_on_off_settings(doc, 'autoHyphenation', bool(document_settings['hyphenation']))

    defaults = spec.get('defaults', {})
    normal = doc.styles['Normal']
    _apply_font(normal, defaults, fonts=fonts, colors=colors)
    _apply_paragraph_format(normal, defaults.get('paragraph', defaults))

    for style_spec in spec.get('styles', []):
        _apply_style_spec(doc, style_spec, fonts=fonts, colors=colors)
    default_paragraph = defaults.get('paragraph', defaults)
    for style in doc.styles:
        if defaults.get('language') is not None:
            _set_style_language(style, str(defaults['language']))
        if style.type != WD_STYLE_TYPE.PARAGRAPH:
            continue
        if default_paragraph.get('widow_control') is not None:
            _set_on_off_ppr(style, 'widowControl', bool(default_paragraph['widow_control']))
        if default_paragraph.get('snap_to_grid') is not None:
            _set_snap_to_grid(style, bool(default_paragraph['snap_to_grid']))

    header_text = spec.get('header', {}).get('text', '')
    header = section.header
    hdr_p = header.paragraphs[0]
    hdr_p.text = str(header_text)
    hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    if spec.get('footer', {}).get('page_number', True):
        fld_simple = OxmlElement('w:fldSimple')
        fld_simple.set(qn('w:instr'), 'PAGE')
        run_elem = OxmlElement('w:r')
        run_elem.append(fld_simple)
        f_p._p.append(run_elem)

    sample = spec.get('sample', {})
    if sample.get('include', True):
        title = sample.get('title', 'Reference DOCX Style Specimen')
        doc.add_heading(str(title), level=1)
        doc.add_paragraph(
            'This specimen page keeps custom styles present in the reference DOCX. '
            'Remove it from final generated documents if it appears in output.'
        )
        for style_spec in spec.get('styles', []):
            if style_spec.get('type', 'paragraph') != 'paragraph':
                continue
            p = doc.add_paragraph(str(style_spec.get('sample_text') or style_spec['name']))
            p.style = doc.styles[style_spec['name']]
        table_styles = [s for s in spec.get('styles', []) if s.get('type') == 'table']
        for style_spec in table_styles:
            doc.add_paragraph(str(style_spec.get('name')))
            table = doc.add_table(rows=2, cols=2)
            table.style = doc.styles[style_spec['name']]
            table.cell(0, 0).text = 'Header 1'
            table.cell(0, 1).text = 'Header 2'
            table.cell(1, 0).text = 'Sample'
            table.cell(1, 1).text = 'Sample'

    doc.save(out_path)


def create_reference(path: str):
    """Create a reference DOCX containing the project's base styles.

    The generated file can be used with Pandoc's `--reference-doc` option
    to control Word styles for converted documents.
    """
    doc = Document()
    section = doc.sections[0]
    # A4 page size
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    # Use neutral manuscript margins to better match markdown preview density.
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)

    styles = doc.styles

    # VS Code markdown preview settings:
    # - markdown.preview.fontFamily defaults to a system sans stack (Segoe UI on Windows)
    # - markdown.preview.fontSize is set to 12 in this environment
    # - markdown.preview.lineHeight default is 1.6
    body_font = 'Segoe UI'
    heading_font = 'Segoe UI Semibold'
    caption_font = 'Segoe UI'
    code_font = 'Consolas'
    base_font_size_pt = 12.0
    base_line_height = 1.6

    # Normal (body) style
    normal = styles['Normal']
    normal.font.name = body_font
    normal.font.size = Pt(base_font_size_pt)
    pformat = normal.paragraph_format
    pformat.space_before = Pt(0)
    pformat.space_after = Pt(12)
    pformat.line_spacing = base_line_height
    pformat.first_line_indent = Mm(0)
    _set_style_rfonts(normal, body_font)

    # Match markdown.css heading scale:
    # h1:2em, h2:1.5em, h3:1.25em, h4:1em, h5:0.875em, h6:0.85em
    heading_specs = (
        ('Heading 1', 24.0, 0.0, 12.0, True),
        ('Heading 2', 18.0, 18.0, 12.0, True),
        ('Heading 3', 15.0, 18.0, 12.0, False),
        ('Heading 4', 12.0, 18.0, 12.0, False),
        ('Heading 5', 10.5, 18.0, 12.0, False),
        ('Heading 6', 10.2, 18.0, 12.0, False),
    )
    for name, size_pt, before_pt, after_pt, with_bottom_rule in heading_specs:
        if name not in styles:
            continue
        heading = styles[name]
        heading.font.name = heading_font
        heading.font.size = Pt(size_pt)
        heading.font.bold = True
        heading.font.color.rgb = RGBColor(0, 0, 0)
        heading.paragraph_format.space_before = Pt(before_pt)
        heading.paragraph_format.space_after = Pt(after_pt)
        heading.paragraph_format.line_spacing = 1.25
        heading.paragraph_format.page_break_before = False
        _set_style_rfonts(heading, heading_font)
        if with_bottom_rule:
            _set_p_borders(
                heading,
                {
                    'bottom': {'val': 'single', 'sz': '6', 'space': '2', 'color': 'D0D7DE'},
                },
            )
        else:
            _set_p_borders(heading, None)

    # Caption style
    if 'Caption' not in styles:
        caption = styles.add_style('Caption', WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles['Caption']
    caption.font.name = caption_font
    caption.font.size = Pt(10)
    caption.font.italic = False
    caption.font.color.rgb = RGBColor(0x57, 0x57, 0x57)
    caption.paragraph_format.space_before = Pt(6)
    caption.paragraph_format.space_after = Pt(6)
    caption.paragraph_format.line_spacing = base_line_height
    _set_style_rfonts(caption, caption_font)

    # Block Text (used for markdown blockquotes)
    if 'Block Text' not in styles:
        block_text = styles.add_style('Block Text', WD_STYLE_TYPE.PARAGRAPH)
    else:
        block_text = styles['Block Text']
    block_text.base_style = styles['Normal']
    block_text.font.name = body_font
    _set_style_rfonts(block_text, body_font)
    block_text.paragraph_format.left_indent = Pt(10)
    block_text.paragraph_format.right_indent = Pt(16)
    block_text.paragraph_format.space_before = Pt(0)
    block_text.paragraph_format.space_after = Pt(12)
    block_text.paragraph_format.line_spacing = base_line_height
    _set_p_borders(
        block_text,
        {
            'left': {'val': 'single', 'sz': '30', 'space': '8', 'color': 'D0D7DE'},
        },
    )
    _set_shading(block_text, fill=None)

    # Intense Quote should follow the same subdued blockquote styling.
    if 'Intense Quote' not in styles:
        bq = styles.add_style('Intense Quote', WD_STYLE_TYPE.PARAGRAPH)
    else:
        bq = styles['Intense Quote']
    bq.base_style = block_text
    _set_style_rfonts(bq, body_font)
    _set_p_borders(
        bq,
        {
            'left': {'val': 'single', 'sz': '30', 'space': '8', 'color': 'D0D7DE'},
        },
    )
    _set_shading(bq, fill=None)

    # List styles inherit the same body typography as markdown preview.
    if 'List Paragraph' not in styles:
        list_style = styles.add_style('List Paragraph', WD_STYLE_TYPE.PARAGRAPH)
    else:
        list_style = styles['List Paragraph']
    list_style.base_style = styles['Normal']
    list_style.font.name = body_font
    list_style.font.italic = False
    list_style.paragraph_format.line_spacing = base_line_height
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(8.4)
    _set_style_rfonts(list_style, body_font)
    _add_contextual_spacing(list_style)

    # List Bullet 2
    if 'List Bullet 2' not in styles and 'ListBullet2' not in styles:
        list_bullet2 = styles.add_style('List Bullet 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        list_bullet2 = styles['List Bullet 2'] if 'List Bullet 2' in styles else styles['ListBullet2']
    list_bullet2.base_style = styles['Normal']
    list_bullet2.paragraph_format.line_spacing = base_line_height
    list_bullet2.paragraph_format.space_before = Pt(0)
    list_bullet2.paragraph_format.space_after = Pt(8.4)
    _add_contextual_spacing(list_bullet2)

    # List Number 2
    if 'List Number 2' not in styles and 'ListNumber2' not in styles:
        list_number2 = styles.add_style('List Number 2', WD_STYLE_TYPE.PARAGRAPH)
    else:
        list_number2 = styles['List Number 2'] if 'List Number 2' in styles else styles['ListNumber2']
    list_number2.base_style = styles['Normal']
    list_number2.paragraph_format.line_spacing = base_line_height
    list_number2.paragraph_format.space_before = Pt(0)
    list_number2.paragraph_format.space_after = Pt(8.4)
    _add_contextual_spacing(list_number2)

    # Monospace / code style
    if 'Code' not in styles:
        code_style = styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)
    else:
        code_style = styles['Code']
    code_style.base_style = styles['Normal']
    code_style.font.name = code_font
    code_style.font.size = Pt(base_font_size_pt)
    code_style.paragraph_format.space_before = Pt(12)
    code_style.paragraph_format.space_after = Pt(12)
    code_style.paragraph_format.line_spacing = 1.357
    _set_style_rfonts(code_style, code_font)
    _set_shading(code_style, fill='F6F8FA')
    _set_p_borders(
        code_style,
        {
            'top': {'val': 'single', 'sz': '6', 'space': '0', 'color': 'D0D7DE'},
            'left': {'val': 'single', 'sz': '6', 'space': '0', 'color': 'D0D7DE'},
            'bottom': {'val': 'single', 'sz': '6', 'space': '0', 'color': 'D0D7DE'},
            'right': {'val': 'single', 'sz': '6', 'space': '0', 'color': 'D0D7DE'},
        },
    )

    # Header placeholder (document title, editable in Word)
    header = section.header
    hdr_p = header.paragraphs[0]
    hdr_p.text = ''
    hdr_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hdr_p.style = styles['Normal']

    # Footer with centered page number field
    footer = section.footer
    f_p = footer.paragraphs[0]
    f_p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # insert PAGE field
    fld_simple = OxmlElement('w:fldSimple')
    fld_simple.set(qn('w:instr'), 'PAGE')
    run_elem = OxmlElement('w:r')
    run_elem.append(fld_simple)
    f_p._p.append(run_elem)

    doc.add_paragraph('Reference docx for styles — remove this placeholder page before use.')
    doc.save(path)


def _load_create_reference_docx():
    """
    Resolve create_reference_docx in both invocation modes:
    - package mode: python -m scripts ...
    - script mode:  python scripts/generate_reference_docx.py ...
    """
    try:
        from .docx_to_markdown import create_reference_docx  # type: ignore[reportMissingImports]
    except ImportError:
        from scripts.docx_to_markdown import create_reference_docx  # type: ignore[reportMissingImports]
    return create_reference_docx


def _build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        '--input',
        default=None,
        help='Optional source DOCX to extract styles from. When omitted, built-in defaults are used.',
    )
    parser.add_argument(
        '--spec',
        default=None,
        help='Optional YAML style specification for generating a reference DOCX.',
    )
    parser.add_argument('--out', default='reference.docx', help='Output path for reference DOCX')
    parser.add_argument(
        '--preserve-headers',
        action='store_true',
        help='When using --input, keep header/footer parts from the source DOCX.',
    )
    return parser


def main(argv: list[str] | None = None) -> int:
    parser = _build_parser()
    args = parser.parse_args(argv)
    out_path = Path(args.out).expanduser()
    out_path.parent.mkdir(parents=True, exist_ok=True)

    if args.input and args.spec:
        print('Error: use either --input or --spec, not both.', file=sys.stderr)
        return 1

    if args.spec:
        spec_path = Path(args.spec).expanduser().resolve()
        if not spec_path.exists():
            print(f'Style spec not found: {spec_path}', file=sys.stderr)
            return 1
        if not spec_path.is_file():
            print(f'Style spec path must be a file: {spec_path}', file=sys.stderr)
            return 1
        try:
            create_reference_from_spec(str(spec_path), str(out_path.resolve()))
        except Exception as exc:
            print(f'Error generating reference DOCX from spec: {exc}', file=sys.stderr)
            return 1
        print(f'Wrote reference docx to {out_path} (source spec: {spec_path})')
        return 0

    if args.input:
        source_path = Path(args.input).expanduser().resolve()
        if not source_path.exists():
            print(f'Input DOCX not found: {source_path}', file=sys.stderr)
            return 1
        if not source_path.is_file():
            print(f'Input path must be a file: {source_path}', file=sys.stderr)
            return 1
        if source_path.suffix.lower() != '.docx':
            print(f'Input file must have .docx extension: {source_path}', file=sys.stderr)
            return 1

        create_reference_docx = _load_create_reference_docx()
        create_reference_docx(source_path, out_path.resolve(), keep_headers=args.preserve_headers)
        print(f'Wrote reference docx to {out_path} (source styles from {source_path})')
        return 0

    if args.preserve_headers:
        print('Warning: --preserve-headers is ignored unless --input is provided.', file=sys.stderr)

    create_reference(str(out_path))
    print(f'Wrote reference docx to {out_path}')
    return 0


if __name__ == '__main__':
    sys.exit(main())

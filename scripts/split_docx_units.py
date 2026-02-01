import argparse
import os
import re
import unicodedata
from copy import deepcopy

import docx
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def iter_block_items(parent):
    for child in parent.element.body.iterchildren():
        if child.tag == qn('w:p'):
            yield Paragraph(child, parent)
        elif child.tag == qn('w:tbl'):
            yield Table(child, parent)


def block_text(block):
    if isinstance(block, Paragraph):
        return block.text.strip()
    texts = []
    for row in block.rows:
        for cell in row.cells:
            cell_text = cell.text.strip()
            if cell_text:
                texts.append(cell_text)
    return "\n".join(texts)


def slugify(title: str) -> str:
    normalized = unicodedata.normalize("NFKD", title)
    ascii_title = normalized.encode("ascii", "ignore").decode("ascii")
    slug = re.sub(r'[^A-Za-z0-9]+', '-', ascii_title).strip('-').lower()
    if not slug:
        return 'section'
    return slug[:60]


def _is_heading_level(paragraph, level: int) -> bool:
    try:
        style = paragraph.style
    except Exception:
        return False
    if not style:
        return False
    name = getattr(style, "name", "") or ""
    style_id = getattr(style, "style_id", "") or ""
    if name.strip().lower() == f"heading {level}".lower():
        return True
    if style_id.replace(" ", "").lower() == f"heading{level}".lower():
        return True
    return False


def find_units(blocks, heading_level: int):
    unit_starts = []
    for idx, (block, text) in enumerate(blocks):
        if isinstance(block, Paragraph) and _is_heading_level(block, heading_level):
            unit_number = str(len(unit_starts) + 1)
            unit_title = text.strip() or f"Unit {unit_number}"
            unit_starts.append((idx, unit_number, unit_title))
    return unit_starts


def remove_images(doc_to_clean):
    for element in doc_to_clean.element.body.iter():
        for child in list(element):
            if child.tag in {qn("w:drawing"), qn("w:pict")}:
                element.remove(child)


def split_docx(
    input_path: str,
    output_dir: str,
    ext_dir: str,
    heading_level: int,
    include_front_matter: bool,
) -> list[tuple[str, str]]:
    doc = docx.Document(input_path)
    blocks = [(block, block_text(block)) for block in iter_block_items(doc)]
    unit_starts = find_units(blocks, heading_level)
    if not unit_starts:
        raise ValueError(f"No unit headings found in {input_path}")

    unit_ranges = []
    for i, (start_idx, unit_num, unit_title) in enumerate(unit_starts):
        end_idx = unit_starts[i + 1][0] if i + 1 < len(unit_starts) else len(blocks)
        unit_ranges.append((unit_num, unit_title, start_idx, end_idx))

    front_range = None
    first_unit_num, first_unit_title, first_start, first_end = unit_ranges[0]
    if first_start > 0:
        if include_front_matter:
            front_range = (0, first_start)
        else:
            unit_ranges[0] = (first_unit_num, first_unit_title, 0, first_end)

    os.makedirs(ext_dir, exist_ok=True)

    created = []
    if front_range:
        start_idx, end_idx = front_range
        front_doc = docx.Document()
        if front_doc.paragraphs:
            p = front_doc.paragraphs[0]._element
            p.getparent().remove(p)
        for block, _text in blocks[start_idx:end_idx]:
            front_doc.element.body.append(deepcopy(block._element))
        remove_images(front_doc)
        front_path = os.path.join(ext_dir, "00-front-matter.docx")
        front_doc.save(front_path)
        created.append(("00", "front-matter"))
    for unit_num, unit_title, start_idx, end_idx in unit_ranges:
        safe_title = slugify(unit_title)
        unit_doc = docx.Document()
        if unit_doc.paragraphs:
            p = unit_doc.paragraphs[0]._element
            p.getparent().remove(p)
        for block, _text in blocks[start_idx:end_idx]:
            unit_doc.element.body.append(deepcopy(block._element))
        remove_images(unit_doc)
        out_path = os.path.join(ext_dir, f"{int(unit_num):02d}-{safe_title}.docx")
        unit_doc.save(out_path)
        created.append((unit_num, safe_title))
    return created


def main():
    parser = argparse.ArgumentParser(
        description="Split a DOCX into unit-level DOCX files, removing images."
    )
    parser.add_argument("input", help="Path to the source DOCX")
    parser.add_argument(
        "output",
        nargs="?",
        default=None,
        help='Output directory (default: "<input folder>\\out")',
    )
    parser.add_argument(
        "--unit-heading-level",
        type=int,
        default=1,
        help="Heading level that marks the start of a unit (default: 1).",
    )
    parser.add_argument(
        "--no-front-matter",
        action="store_true",
        help="Include front matter in the first unit instead of a separate file.",
    )
    args = parser.parse_args()

    input_path = os.path.abspath(args.input)
    input_dir = os.path.dirname(input_path)
    input_stem, input_ext = os.path.splitext(os.path.basename(input_path))
    parent_name = os.path.basename(input_dir)
    desired_parent = os.path.join(input_dir, input_stem) if parent_name != input_stem else input_dir
    if not os.path.isdir(desired_parent):
        os.makedirs(desired_parent, exist_ok=True)
    output_dir = os.path.abspath(
        args.output if args.output else os.path.join(desired_parent, "out")
    )
    ext_dir = os.path.join(output_dir, input_ext)

    created = split_docx(
        input_path,
        output_dir,
        ext_dir,
        args.unit_heading_level,
        include_front_matter=not args.no_front_matter,
    )
    for unit_num, title in created:
        if unit_num == "00":
            print(f"{unit_num}-front-matter")
        else:
            print(f"{int(unit_num):02d}-{slugify(title)}")

    # Move the source DOCX into the same-named folder last, after all processing.
    if parent_name != input_stem:
        moved_path = os.path.join(desired_parent, os.path.basename(input_path))
        if os.path.normcase(moved_path) != os.path.normcase(input_path):
            os.replace(input_path, moved_path)


if __name__ == "__main__":
    main()
